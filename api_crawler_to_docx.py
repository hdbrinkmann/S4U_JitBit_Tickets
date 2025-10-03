#!/usr/bin/env python3
"""
API Documentation Crawler to Word Document Generator

This script specifically crawls API documentation websites, extracts API functions
from the main page, and creates a well-organized Word document with each API
as its own chapter containing both overview and detailed information.
"""

import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
import logging
import argparse
import re
from datetime import datetime
from collections import OrderedDict


class ApiDocumentationCrawler:
    def __init__(self, start_url, delay=1):
        """
        Initialize the API documentation crawler.
        
        Args:
            start_url (str): The starting URL to crawl
            delay (float): Delay between requests in seconds (default: 1)
        """
        self.start_url = start_url
        self.base_domain = urlparse(start_url).netloc
        self.delay = delay
        self.apis = OrderedDict()  # Store API information
        
        # Setup logging
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )
        self.logger = logging.getLogger(__name__)
        
        # Setup session with headers
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })

    def clean_text(self, text):
        """Clean and normalize text content."""
        text = re.sub(r'\s+', ' ', text.strip())
        return text

    def extract_api_functions_from_main_page(self, soup):
        """Extract API functions and their descriptions from the main page."""
        apis = OrderedDict()
        
        # Look for the main content area that lists API functions
        # This typically contains links to individual API operations
        api_links = soup.find_all('a', href=True)
        
        for link in api_links:
            href = link.get('href', '')
            if '?op=' in href:
                # Extract API name from the operation parameter
                api_name = href.split('?op=')[-1]
                api_url = urljoin(self.start_url, href)
                
                # Get the link text and surrounding context for description
                link_text = self.clean_text(link.get_text())
                
                # Try to find description by looking at the surrounding elements
                description = ""
                
                # Look for description in parent elements
                parent = link.parent
                if parent:
                    # Check if there's descriptive text near the link
                    parent_text = self.clean_text(parent.get_text())
                    if parent_text and len(parent_text) > len(link_text):
                        # Extract the part that's not the link text
                        description = parent_text.replace(link_text, '').strip()
                        # Clean up common patterns
                        description = re.sub(r'^[-\s]*', '', description)
                        description = re.sub(r'[-\s]*$', '', description)
                
                # Store API information
                apis[api_name] = {
                    'name': api_name,
                    'display_name': link_text or api_name.replace('_', ' ').title(),
                    'url': api_url,
                    'main_page_description': description,
                    'detailed_content': None
                }
        
        return apis

    def fetch_detailed_api_content(self, api_info):
        """Fetch detailed content for a specific API."""
        try:
            self.logger.info(f"Fetching details for API: {api_info['name']}")
            response = self.session.get(api_info['url'], timeout=30)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()
            
            # Extract meaningful content
            content_elements = []
            
            # Find the main content area
            main_content = soup.find('body')
            if main_content:
                for element in main_content.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li', 'div', 'pre', 'code']):
                    text = self.clean_text(element.get_text())
                    if text and len(text) > 5:  # Include substantial text
                        element_type = element.name
                        content_elements.append({
                            'type': element_type,
                            'text': text,
                            'level': int(element_type[1]) if element_type.startswith('h') else 0
                        })
            
            api_info['detailed_content'] = content_elements
            time.sleep(self.delay)  # Be respectful
            
        except Exception as e:
            self.logger.error(f"Error fetching details for {api_info['name']}: {e}")
            api_info['detailed_content'] = []

    def crawl_api_documentation(self):
        """Crawl the API documentation starting from the main page."""
        self.logger.info(f"Starting to crawl API documentation: {self.start_url}")
        
        try:
            # Fetch main page
            response = self.session.get(self.start_url, timeout=30)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Extract API functions from main page
            self.apis = self.extract_api_functions_from_main_page(soup)
            
            self.logger.info(f"Found {len(self.apis)} API functions")
            
            # Fetch detailed content for each API
            for api_name, api_info in self.apis.items():
                self.fetch_detailed_api_content(api_info)
            
            self.logger.info("API documentation crawling completed")
            
        except Exception as e:
            self.logger.error(f"Error crawling main page: {e}")
            raise

    def create_word_document(self, output_filename):
        """Create a well-formatted Word document with the API documentation."""
        self.logger.info("Creating Word document...")
        
        doc = Document()
        
        # Set up document styles
        self.setup_document_styles(doc)
        
        # Add title
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(f"API Documentation: {self.base_domain}")
        title_run.font.size = Pt(18)
        title_run.bold = True
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Source URL: {self.start_url}")
        doc.add_paragraph(f"Total API functions: {len(self.apis)}")
        doc.add_paragraph("")
        
        # Add table of contents
        doc.add_heading("Table of Contents", level=1)
        for i, (api_name, api_info) in enumerate(self.apis.items(), 1):
            toc_paragraph = doc.add_paragraph(f"{i}. {api_info['display_name']}")
            toc_paragraph.style = 'List Number'
        doc.add_page_break()
        
        # Add content for each API as chapters
        for i, (api_name, api_info) in enumerate(self.apis.items(), 1):
            # Chapter header
            doc.add_heading(api_info['display_name'], level=1)
            
            # Add API URL as subtitle
            url_paragraph = doc.add_paragraph()
            url_run = url_paragraph.add_run(f"Endpoint: {api_info['url']}")
            url_run.font.size = Pt(9)
            url_run.italic = True
            doc.add_paragraph("")
            
            # Add main page description if available
            if api_info['main_page_description']:
                doc.add_heading("Overview", level=2)
                doc.add_paragraph(api_info['main_page_description'])
                doc.add_paragraph("")
            
            # Add detailed content
            if api_info['detailed_content']:
                doc.add_heading("Detailed Information", level=2)
                
                for element in api_info['detailed_content']:
                    if element['type'].startswith('h') and element['level'] > 0:
                        # Adjust heading level to fit under the API chapter
                        adjusted_level = min(element['level'] + 2, 6)
                        doc.add_heading(element['text'], level=adjusted_level)
                    elif element['type'] == 'p':
                        doc.add_paragraph(element['text'])
                    elif element['type'] == 'li':
                        doc.add_paragraph(element['text'], style='List Bullet')
                    elif element['type'] in ['pre', 'code']:
                        # Format code blocks differently
                        code_para = doc.add_paragraph(element['text'])
                        code_para.style = 'Normal'
                        for run in code_para.runs:
                            run.font.name = 'Consolas'
                            run.font.size = Pt(10)
                    else:
                        doc.add_paragraph(element['text'])
            else:
                doc.add_paragraph("No detailed information available.")
            
            # Add page break between chapters (except for the last one)
            if i < len(self.apis):
                doc.add_page_break()
        
        # Save document
        doc.save(output_filename)
        self.logger.info(f"Word document saved as: {output_filename}")

    def setup_document_styles(self, doc):
        """Setup custom styles for the document."""
        styles = doc.styles
        
        try:
            normal_style = styles['Normal']
            normal_style.font.size = Pt(11)
            normal_style.font.name = 'Calibri'
        except:
            pass


def main():
    """Main function to run the API documentation crawler."""
    parser = argparse.ArgumentParser(description='Crawl API documentation and generate a Word document')
    parser.add_argument('url', help='The starting URL to crawl')
    parser.add_argument('-o', '--output', default='api_documentation.docx',
                       help='Output Word document filename (default: api_documentation.docx)')
    parser.add_argument('--delay', type=float, default=1.0,
                       help='Delay between requests in seconds (default: 1.0)')
    
    args = parser.parse_args()
    
    # Create crawler instance
    crawler = ApiDocumentationCrawler(
        start_url=args.url,
        delay=args.delay
    )
    
    try:
        # Crawl the API documentation
        crawler.crawl_api_documentation()
        
        # Create Word document
        crawler.create_word_document(args.output)
        
        print(f"\n✅ Successfully created Word document: {args.output}")
        print(f"📊 Total API functions documented: {len(crawler.apis)}")
        
        # Show API list
        print("\n📋 API Functions Found:")
        for i, (api_name, api_info) in enumerate(crawler.apis.items(), 1):
            print(f"  {i}. {api_info['display_name']}")
        
    except KeyboardInterrupt:
        print("\n⚠️  Crawling interrupted by user")
    except Exception as e:
        print(f"\n❌ Error: {e}")


if __name__ == "__main__":
    main()
