#!/usr/bin/env python3
"""
Generate a single DOCX from a Jitbit Knowledgebase JSON export
with an API-first approach for loading images.

What this script does:
- Renders one article per page with subject + metadata
- Converts Body HTML into simple paragraphs/lists/pre/tables/images
- Loads all Jitbit images/attachments via the official API:
    GET {base}/helpdesk/api/attachment?id={FileID}
  using a Bearer token from the environment (.env: JITBIT_API_TOKEN=...)

- External images (e.g., imgur, teams CDN) are fetched directly without cookies

Usage:
  1) Ensure .env contains:
       JITBIT_API_TOKEN=...
       (optional) JITBIT_BASE_URL=https://support.example.com
     Note: If JITBIT_BASE_URL is not set, the script derives the API base
     from export_info.api_base_url in the JSON.

  2) Run:
     python kb_to_pdf.py --input JitBit_Knowledgebase.json --output Knowledgebase.docx --verbose true
"""

import argparse
import io
import json
import os
import sys
from typing import Dict, List, Optional, Tuple, Set
from urllib.parse import urljoin, urlparse, parse_qs
import re

try:
    from dotenv import load_dotenv
    load_dotenv()
except Exception:
    pass

import requests
from bs4 import BeautifulSoup, NavigableString, Tag
from PIL import Image as PILImage

# python-docx
from docx import Document
from docx.shared import Inches, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def str2bool(v: str) -> bool:
    return str(v).lower() in {"1", "true", "t", "yes", "y"}


def xml_safe(s: Optional[str]) -> str:
    """
    Remove XML 1.0 illegal control characters (except TAB, LF, CR) that cause python-docx to fail.
    """
    if not s:
        return ""
    return re.sub(r"[\x00-\x08\x0B-\x0C\x0E-\x1F]", "", s)


def truncate_text(s: str, max_len: int = 300) -> str:
    return s if len(s) <= max_len else s[: max_len - 1] + "…"


def sanitize_url(raw: Optional[str]) -> Optional[str]:
    if raw is None:
        return None
    s = str(raw).strip()
    if not s:
        return None
    # Strip surrounding quotes
    if (s.startswith('"') and s.endswith('"')) or (s.startswith("'") and s.endswith("'")):
        s = s[1:-1]
    # Cut at first whitespace
    s = s.split()[0]
    # Remove any embedded tags
    for ch in ("<", ">"):
        pos = s.find(ch)
        if pos != -1:
            s = s[:pos]
    return s or None


def resolve_url(src: str, base: Optional[str]) -> Optional[str]:
    if not src:
        return None
    src = src.strip()
    if src.startswith("http://") or src.startswith("https://"):
        return src
    if base:
        return urljoin(base.rstrip("/") + "/", src.lstrip("/"))
    return None


def extract_file_id_from_url(url_str: str) -> Optional[str]:
    """
    Extracts a numeric FileID from typical Jitbit URLs:
      - /helpdesk/File/Get/26355
      - /helpdesk/File/Get?id=26355
    Returns None if no numeric id present.
    """
    try:
        pu = urlparse(url_str)
        # Query ?id=123
        qs = parse_qs(pu.query or "")
        if "id" in qs:
            v = qs["id"][0]
            if str(v).isdigit():
                return str(v)
        # Path segment
        segs = [s for s in (pu.path or "").split("/") if s]
        if segs:
            last = segs[-1]
            if last.isdigit():
                return last
    except Exception:
        pass
    return None


def derive_api_root(api_base_url: Optional[str], env_base_url: Optional[str], verbose: bool = False) -> Optional[str]:
    """
    Derive the API root:
      If api_base_url path contains '/helpdesk' -> {scheme}://{host}/helpdesk/api
      Else -> {scheme}://{host}/api
    Falls back to env_base_url if api_base_url is missing.
    """
    base = api_base_url or env_base_url
    if not base:
        return None
    try:
        pu = urlparse(base)
        scheme = pu.scheme or "https"
        netloc = pu.netloc
        path = (pu.path or "").lower()
        if not netloc:
            return None
        if "/helpdesk" in path:
            root = f"{scheme}://{netloc}/helpdesk/api"
        else:
            root = f"{scheme}://{netloc}/api"
        if verbose:
            print(f"[INFO] API root derived: {root}", file=sys.stderr)
        return root
    except Exception:
        return None


class JitbitFetcher:
    def __init__(self, api_root: Optional[str], token: Optional[str], timeout: float = 15.0, verbose: bool = False):
        self.api_root = api_root
        self.token = token
        self.timeout = timeout
        self.verbose = verbose

        self.session = requests.Session()
        self.session.headers.update({
            "User-Agent": "KB-DOCX/1.0",
            "Accept": "*/*",
        })
        if token:
            self.session.headers["Authorization"] = f"Bearer {token}"

    def _warn(self, msg: str):
        if self.verbose:
            print(msg, file=sys.stderr)

    def fetch_attachment_by_id(self, file_id: str) -> Optional[bytes]:
        """
        Fetches an attachment by FileID trying the correct API root(s).
        Many Jitbit installs live under '/helpdesk', so we try:
          1) {scheme}://{host}/helpdesk/api/attachment?id=...
          2) {scheme}://{host}/api/attachment?id=...
        We prioritize any '/helpdesk' root if self.api_root includes it.
        """
        if not self.api_root or not self.token:
            self._warn(f"[WARN] Missing API root or token for Jitbit attachment id={file_id}")
            return None

        try:
            p = urlparse(self.api_root)
            scheme = p.scheme or "https"
            netloc = p.netloc
            path = (p.path or "").lower()
            if not netloc:
                self._warn(f"[WARN] Invalid api_root (no host): {self.api_root}")
                return None

            candidates = []
            # Prioritize '/helpdesk/api' if api_root path mentions helpdesk
            if "/helpdesk" in path:
                candidates.append(f"{scheme}://{netloc}/helpdesk/api")
            # Always try '/helpdesk/api' first for safety if not already present
            if f"{scheme}://{netloc}/helpdesk/api" not in candidates:
                candidates.append(f"{scheme}://{netloc}/helpdesk/api")
            # Then plain '/api'
            candidates.append(f"{scheme}://{netloc}/api")

            last_err: Optional[Exception] = None
            for root in candidates:
                url = f"{root.rstrip('/')}/attachment?id={file_id}"
                try:
                    self._warn(f"[INFO] GET {url}")
                    r = self.session.get(url, timeout=self.timeout)
                    r.raise_for_status()
                    return r.content
                except Exception as e:
                    last_err = e
                    self._warn(f"[INFO] Candidate failed: {e}")
                    continue

            if last_err:
                self._warn(f"[WARN] API fetch failed for id={file_id}: {last_err}")
            return None
        except Exception as e:
            self._warn(f"[WARN] API fetch failed for id={file_id}: {e}")
            return None

    def fetch_generic_image(self, url: str) -> Optional[bytes]:
        try:
            self._warn(f"[INFO] GET (external) {url}")
            r = self.session.get(url, timeout=self.timeout)
            r.raise_for_status()
            # Validate image
            with PILImage.open(io.BytesIO(r.content)) as _:
                pass
            return r.content
        except Exception as e:
            self._warn(f"[WARN] External image fetch failed for {url}: {e}")
            return None


# ---- DOCX helpers ----

A4_INCH = (8.27, 11.69)
LETTER_INCH = (8.5, 11.0)


def set_page_size_and_margins(doc: Document, page: str, margin_pt: float) -> None:
    """
    Configure the document's first section to the requested page size and margins.
    margin_pt is in points (1/72 inch).
    """
    section = doc.sections[0]

    if page.upper() == "A4":
        w_in, h_in = A4_INCH
    else:
        w_in, h_in = LETTER_INCH

    section.page_width = Inches(w_in)
    section.page_height = Inches(h_in)

    m_in = float(margin_pt) / 72.0
    section.left_margin = Inches(m_in)
    section.right_margin = Inches(m_in)
    section.top_margin = Inches(m_in)
    section.bottom_margin = Inches(m_in)


def get_usable_emu(doc: Document) -> Tuple[int, int]:
    """
    Return (usable_width_emu, usable_height_emu) for the first section.
    """
    s = doc.sections[0]
    usable_w = int(s.page_width - s.left_margin - s.right_margin)
    usable_h = int(s.page_height - s.top_margin - s.bottom_margin)
    return usable_w, usable_h


def _bytes_to_image_dims_emu(img_bytes: bytes) -> Optional[Tuple[int, int]]:
    """
    Returns (width_emu, height_emu) using image DPI metadata if available, else assumes 96 DPI.
    """
    try:
        with PILImage.open(io.BytesIO(img_bytes)) as im:
            w_px, h_px = im.width, im.height
            dpi = im.info.get("dpi", (96, 96))
            dpi_x = float(dpi[0] or 96.0)
            dpi_y = float(dpi[1] or 96.0)
            w_in = w_px / dpi_x
            h_in = h_px / dpi_y
            return int(Emu(Inches(w_in))), int(Emu(Inches(h_in)))
    except Exception:
        return None


def add_image_placeholder_docx(doc: Document, url: str, label: Optional[str] = None, auth_hint: bool = False):
    safe_href = xml_safe(url or "")
    display_text = truncate_text(safe_href, 300)
    label_text = f"{label} – " if label else ""
    hint_text = " (evtl. Anmeldung/Cookies erforderlich)" if auth_hint else ""
    p = doc.add_paragraph()
    run = p.add_run(xml_safe(f"{label_text}Bild konnte nicht geladen werden: {display_text}{hint_text}"))
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def add_image_from_src_docx(
    doc: Document,
    src: str,
    kb_base_url: Optional[str],
    fetcher: JitbitFetcher,
    include_images: bool,
    max_width_emu: int,
    max_height_emu: int,
    add_placeholders: bool,
):
    if not include_images:
        return
    clean = sanitize_url(src)
    if not clean:
        return
    abs_url = resolve_url(clean, kb_base_url) or clean

    # Try to extract a Jitbit FileID
    fid = extract_file_id_from_url(abs_url)
    data = None
    if fid:
        data = fetcher.fetch_attachment_by_id(fid)
    if not data:
        # Fallback to generic external download if not a Jitbit file or API fails
        if abs_url and (abs_url.startswith("http://") or abs_url.startswith("https://")):
            data = fetcher.fetch_generic_image(abs_url)

    if not data:
        if add_placeholders:
            add_image_placeholder_docx(doc, abs_url, auth_hint=bool(fid))
        return

    dims = _bytes_to_image_dims_emu(data)
    if not dims:
        if add_placeholders:
            add_image_placeholder_docx(doc, abs_url, auth_hint=bool(fid))
        return

    w_emu, h_emu = dims
    if w_emu <= 0 or h_emu <= 0:
        if add_placeholders:
            add_image_placeholder_docx(doc, abs_url, auth_hint=bool(fid))
        return

    scale = min(max_width_emu / w_emu, max_height_emu / h_emu, 1.0)
    target_w = int(w_emu * scale)
    bio = io.BytesIO(data)
    try:
        doc.add_picture(bio, width=Emu(target_w))
    except Exception:
        if add_placeholders:
            add_image_placeholder_docx(doc, abs_url, auth_hint=bool(fid))


def html_to_docx(
    doc: Document,
    html: str,
    styles: Dict,
    kb_base_url: Optional[str],
    fetcher: JitbitFetcher,
    include_images: bool,
    max_width_emu: int,
    max_height_emu: int,
    add_placeholders: bool,
) -> None:
    if not html:
        return

    soup = BeautifulSoup(html, "html.parser")

    def process_block(node):
        if isinstance(node, NavigableString):
            text = str(node)
            if text and text.strip():
                p = doc.add_paragraph(xml_safe(text))
                return
            return
        if not isinstance(node, Tag):
            return

        name = node.name.lower()

        if name in ("p", "div"):
            # Render text first, then images inside
            node_text_only = BeautifulSoup(str(node), "html.parser")
            for it in node_text_only.find_all("img"):
                it.decompose()
            text = node_text_only.get_text(separator=" ", strip=True)
            if text:
                doc.add_paragraph(xml_safe(text))
            if include_images:
                for img in node.find_all("img"):
                    src = img.get("src")
                    if src:
                        add_image_from_src_docx(
                            doc, src, kb_base_url, fetcher, include_images, max_width_emu, max_height_emu, add_placeholders
                        )

        elif name == "br":
            doc.add_paragraph()

        elif name in ("ul", "ol"):
            ordered = name == "ol"
            for li in node.find_all("li", recursive=False):
                txt = li.get_text(separator=" ", strip=True)
                if txt:
                    p = doc.add_paragraph(xml_safe(txt))
                    p.style = "List Number" if ordered else "List Bullet"

        elif name in ("pre", "code"):
            txt = node.get_text("\n")
            if txt:
                p = doc.add_paragraph()
                run = p.add_run(xml_safe(txt))
                run.font.name = "Courier New"

        elif name == "table":
            # Build a simple table
            rows: List[List[str]] = []
            for tr in node.find_all("tr"):
                cells = tr.find_all(["td", "th"])
                if not cells:
                    continue
                row = [c.get_text(separator=" ", strip=True) for c in cells]
                if any(cell for cell in row):
                    rows.append(row)
            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                # table.style = "Light Grid"  # optional, may not exist on all systems
                for r_idx, r in enumerate(rows):
                    for c_idx, cell in enumerate(r):
                        table.cell(r_idx, c_idx).text = xml_safe(cell)

        elif name == "img":
            src = node.get("src")
            if src:
                add_image_from_src_docx(
                    doc, src, kb_base_url, fetcher, include_images, max_width_emu, max_height_emu, add_placeholders
                )

        else:
            for child in node.children:
                process_block(child)

    for child in soup.contents:
        process_block(child)


def add_attachments_images_docx(
    doc: Document,
    attachments: List[dict],
    kb_base_url: Optional[str],
    fetcher: JitbitFetcher,
    max_width_emu: int,
    max_height_emu: int,
    add_placeholders: bool,
) -> None:
    if not attachments:
        return

    seen: Set[str] = set()

    for att in attachments:
        url = (att.get("Url") or att.get("URL") or "").strip()
        if not url:
            continue
        abs_url = resolve_url(url, kb_base_url) or url
        if abs_url in seen:
            continue
        seen.add(abs_url)

        fid = extract_file_id_from_url(abs_url)
        data = None
        if fid:
            data = fetcher.fetch_attachment_by_id(fid)
        else:
            data = fetcher.fetch_generic_image(abs_url)

        if not data:
            add_image_placeholder_docx(doc, abs_url, label=att.get("FileName") or None, auth_hint=bool(fid))
            continue

        dims = _bytes_to_image_dims_emu(data)
        if not dims:
            add_image_placeholder_docx(doc, abs_url, label=att.get("FileName") or None, auth_hint=bool(fid))
            continue

        w_emu, h_emu = dims
        if w_emu <= 0 or h_emu <= 0:
            add_image_placeholder_docx(doc, abs_url, label=att.get("FileName") or None, auth_hint=bool(fid))
            continue

        scale = min(max_width_emu / w_emu, max_height_emu / h_emu, 1.0)
        target_w = int(w_emu * scale)
        bio = io.BytesIO(data)
        try:
            doc.add_picture(bio, width=Emu(target_w))
        except Exception:
            add_image_placeholder_docx(doc, abs_url, label=att.get("FileName") or None, auth_hint=bool(fid))


def main():
    parser = argparse.ArgumentParser(description="Generate a DOCX from Jitbit Knowledgebase JSON (API-first image fetching).")
    parser.add_argument("--input", "-i", default="JitBit_Knowledgebase.json", help="Path to JSON export file")
    parser.add_argument("--output", "-o", default="Knowledgebase.docx", help="Output DOCX filename")
    parser.add_argument("--page-size", choices=["A4", "LETTER"], default="A4", help="Page size")
    parser.add_argument("--margin", type=float, default=36.0, help="Margins in points (default ~0.5 inch)")
    parser.add_argument("--include-body-images", type=str2bool, default=True, help="Include images from Body HTML")
    parser.add_argument("--include-attachments", type=str2bool, default=True, help="Include images from Attachments")
    parser.add_argument("--timeout", type=float, default=15.0, help="HTTP timeout seconds for API/generic image downloads")
    parser.add_argument("--image-placeholder", type=str2bool, default=True, help="Insert a textual placeholder with link when image download fails")
    parser.add_argument("--attachments-header", type=str2bool, default=False, help="Insert a small heading before the attachments section")
    parser.add_argument("--verbose", type=str2bool, default=False, help="Enable verbose logging")
    args = parser.parse_args()

    # Load JSON
    with open(args.input, "r", encoding="utf-8") as f:
        data = json.load(f)

    export_info = data.get("export_info", {}) or {}
    kb_base_url = export_info.get("api_base_url") or os.getenv("JITBIT_BASE_URL", "")
    kb_base_url = kb_base_url.strip() or None

    api_root = derive_api_root(export_info.get("api_base_url"), os.getenv("JITBIT_BASE_URL", "").strip() or None, verbose=args.verbose)
    token = (os.getenv("JITBIT_API_TOKEN", "") or "").strip()

    if not token:
        print("[WARN] JITBIT_API_TOKEN not set in environment/.env. Jitbit-protected images will not load.", file=sys.stderr)

    fetcher = JitbitFetcher(api_root=api_root, token=token, timeout=args.timeout, verbose=args.verbose)

    # Articles
    articles = data.get("articles", [])
    if not isinstance(articles, list) or not articles:
        print("[INFO] No articles found in input JSON.")
        articles = []

    # Build document
    doc = Document()
    set_page_size_and_margins(doc, args.page_size, args.margin)
    usable_w_emu, usable_h_emu = get_usable_emu(doc)

    for idx, art in enumerate(articles, start=1):
        subject = (art.get("Subject") or "").strip() or "(Ohne Betreff)"
        category = (art.get("CategoryName") or "").strip()
        tagstring = (art.get("TagString") or "").strip()

        # Header
        title_p = doc.add_paragraph(xml_safe(subject))
        title_p.style = "Heading 1"

        # Subheader
        sub_parts = []
        if category:
            sub_parts.append(f"Category: {category}")
        if tagstring:
            sub_parts.append(f"Tags: {tagstring}")
        sub_text = " • ".join(sub_parts) if sub_parts else ""
        if sub_text:
            p = doc.add_paragraph(xml_safe(sub_text))
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        else:
            doc.add_paragraph()

        # Body (API-first image fetching for Jitbit, generic for external)
        body_html = art.get("Body") or art.get("BodyMarkdown") or ""
        html_to_docx(
            doc=doc,
            html=body_html,
            styles={},
            kb_base_url=kb_base_url,
            fetcher=fetcher,
            include_images=args.include_body_images,
            max_width_emu=usable_w_emu,
            max_height_emu=usable_h_emu,
            add_placeholders=args.image_placeholder,
        )

        # Attachments images
        if args.include_attachments:
            att = art.get("Attachments") or []
            if att:
                if args.attachments_header:
                    h = doc.add_paragraph("Anhänge")
                    h.style = "Heading 2"
                add_attachments_images_docx(
                    doc=doc,
                    attachments=att,
                    kb_base_url=kb_base_url,
                    fetcher=fetcher,
                    max_width_emu=usable_w_emu,
                    max_height_emu=usable_h_emu,
                    add_placeholders=args.image_placeholder,
                )

        # Page break
        if idx != len(articles):
            doc.add_page_break()

    doc.save(args.output)
    print(f"[OK] DOCX generated: {args.output}")


if __name__ == "__main__":
    main()
