#!/usr/bin/env python3
"""
Generate Q&A from DOCX files via a two-step pipeline:
1) extract: parse DOCX chapters into an intermediate JSON corpus
2) qa: read chapters JSON and generate Q&A using an OpenAI-compatible endpoint (Scaleway gpt-oss-120b)

Environment variables for LLM:
- SCW_BASE_URL (e.g., https://api.scaleway.ai)  [required for qa]
- SCW_API_KEY  [required for qa]
- SCW_MODEL (default: gpt-oss-120b)
- SCW_JSON_MODE (optional: "1" to request JSON mode if supported)

Examples:
- Extract chapters:
  python scripts/generate_qa_from_docx.py extract --input QA_SOURCE --output QA_CHAPTERS --heading-level 1
- Generate Q&A:
  python scripts/generate_qa_from_docx.py qa --input QA_CHAPTERS --output QA_OUTPUT --max-per-chapter 10
"""

import argparse
import os
import sys
import json
import re
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

# --- .env loader (optional) ---
from pathlib import Path as _Path

def _fallback_load_dotenv() -> None:
    env_path = _Path(".env")
    if not env_path.exists():
        return
    try:
        for line in env_path.read_text(encoding="utf-8").splitlines():
            line = line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, val = line.split("=", 1)
            key = key.strip()
            val = val.strip().strip("'").strip('"')
            if key and key not in os.environ:
                os.environ[key] = val
    except Exception:
        pass

def load_env() -> None:
    try:
        from dotenv import load_dotenv  # type: ignore
        load_dotenv()
    except Exception:
        _fallback_load_dotenv()

# Optional imports (available at runtime if installed)
try:
    import tiktoken  # type: ignore
except Exception:
    tiktoken = None  # type: ignore

try:
    from docx import Document  # type: ignore
except Exception:
    Document = None  # type: ignore

# LLM client (OpenAI-compatible)
OpenAI = None  # lazy import to allow extract without openai installed

# Retry support
try:
    from tenacity import retry, stop_after_attempt, wait_exponential  # type: ignore
except Exception:
    # Fallback no-op decorators if tenacity is not installed
    def retry(*args, **kwargs):
        def wrap(fn):
            return fn
        return wrap

    def stop_after_attempt(*args, **kwargs):
        return None

    def wait_exponential(*args, **kwargs):
        return None


# ----------------------------
# Utilities
# ----------------------------

def list_docx_files(input_dir: str) -> List[str]:
    files = []
    for name in os.listdir(input_dir):
        # Skip Word lock files like "~$..." and hidden files
        if name.startswith("~$") or name.startswith("."):
            continue
        if name.lower().endswith(".docx"):
            path = os.path.join(input_dir, name)
            if os.path.isfile(path):
                files.append(path)
    files.sort()
    return files


def ensure_dir(path: str) -> None:
    os.makedirs(path, exist_ok=True)


def iso_now() -> str:
    return datetime.utcnow().replace(microsecond=0).isoformat() + "Z"


def estimate_tokens(text: str) -> int:
    # Prefer tiktoken if available; else heuristic chars/4
    if tiktoken is not None:
        try:
            enc = tiktoken.get_encoding("cl100k_base")
            return len(enc.encode(text))
        except Exception:
            pass
    # heuristic
    return max(1, len(text) // 4)


def determine_qa_count(token_estimate: int, max_per_chapter: int = 10) -> int:
    # Roughly 350 tokens per Q&A item
    n = round(token_estimate / 350.0)
    n = max(1, min(max_per_chapter, n))
    return n

# --------
# Chunking helpers for large chapters
# --------
def _max_chunk_tokens() -> int:
    v = os.getenv("QA_MAX_CHUNK_TOKENS")
    try:
        return max(1000, int(v)) if v else 3500
    except Exception:
        return 3500


def split_text_by_tokens(text: str, max_tokens: int) -> List[str]:
    """
    Split text into chunks that each approximately stay within max_tokens,
    preferring paragraph boundaries (double newlines).
    """
    if not (text or "").strip():
        return []

    parts = re.split(r"\n{2,}", text)
    chunks: List[str] = []
    cur: List[str] = []
    cur_tokens = 0
    for part in parts:
        t = part.strip()
        if not t:
            continue
        ptoks = estimate_tokens(t)
        if cur and (cur_tokens + ptoks) > max_tokens:
            chunks.append("\n\n".join(cur))
            cur = [t]
            cur_tokens = ptoks
        else:
            cur.append(t)
            cur_tokens += ptoks
    if cur:
        chunks.append("\n\n".join(cur))

    # Merge tiny trailing chunk into previous if too small
    if len(chunks) >= 2 and estimate_tokens(chunks[-1]) < max_tokens * 0.2:
        last = chunks.pop()
        chunks[-1] = chunks[-1] + "\n\n" + last

    return chunks


def compute_chapter_allocations(chapters: List[Dict[str, Any]], max_per_chapter: int, max_per_document: Optional[int]) -> List[int]:
    """
    Compute per-chapter Q&A targets with an optional per-document cap.
    - Base demand per chapter is proportional to token_estimate (≈ tokens/350), clamped to [0, max_per_chapter].
    - If sum(base) <= max_per_document (or no cap), return base.
    - Else, scale proportionally and distribute remainder by largest fractional parts.
    """
    # Build base demand list
    base: List[int] = []
    for ch in chapters:
        te = int(ch.get("token_estimate") or estimate_tokens((ch.get("content") or "")))
        approx = round(te / 350.0)
        base.append(min(max_per_chapter, max(0, approx)))

    if max_per_document is None:
        return base

    total = sum(base)
    if total <= max_per_document:
        return base

    if total == 0:
        return [0 for _ in base]

    # Proportional scaling with remainder distribution
    raw = [b * (max_per_document / float(total)) for b in base]
    floored = [int(x) for x in raw]
    remainder = max_per_document - sum(floored)

    # Distribute +1 to the largest fractional parts
    fracs = sorted(((raw[i] - floored[i], i) for i in range(len(base))), reverse=True)
    i = 0
    while remainder > 0 and i < len(fracs):
        idx = fracs[i][1]
        floored[idx] += 1
        remainder -= 1
        i += 1

    return floored


# ----------------------------
# Coverage-mode helpers
# ----------------------------

def _parse_concepts_json(text: str, max_items: int) -> List[Dict[str, Any]]:
    """
    Expect JSON like:
    { "concepts": [ { "title": "...", "summary": "...", "importance": 1|2|3 } ] }
    """
    try:
        data = json.loads(text)
    except Exception:
        # Try to salvage from a larger blob
        start = text.find("{")
        end = text.rfind("}")
        if start != -1 and end != -1 and end > start:
            try:
                data = json.loads(text[start:end + 1])
            except Exception:
                return []
        else:
            return []
    arr = data.get("concepts")
    if not isinstance(arr, list):
        return []
    out: List[Dict[str, Any]] = []
    for it in arr[:max_items]:
        if not isinstance(it, dict):
            continue
        title = (it.get("title") or "").strip()
        if not title:
            continue
        summary = (it.get("summary") or "").strip()
        importance = it.get("importance")
        try:
            importance = int(importance) if importance is not None else 2
        except Exception:
            importance = 2
        importance = min(3, max(1, importance))
        out.append({"title": title, "summary": summary, "importance": importance})
    return out


def _build_concepts_messages(title: str, content: str, max_items: int) -> List[Dict[str, str]]:
    system = (
        "You are to extract atomic concepts/topics from a chapter. "
        "Capture definitions, procedures, configuration options, exceptions, and important rules. "
        "Keep items concise and distinct."
    )
    user = (
        f"Chapter Title: {title}\n"
        f"Chapter Content:\n{content}\n\n"
        f"Instructions:\n"
        f"- List up to {max_items} distinct concepts.\n"
        "- JSON only, schema:\n"
        '{ "concepts": [ { "title": "string", "summary": "string", "importance": 1|2|3 } ] }\n'
        "- importance: 3=critical, 2=important, 1=minor.\n"
        "- No extra keys or commentary."
    )
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]


def extract_concepts_for_chapter(client, model: str, title: str, content: str, concepts_max: int) -> List[Dict[str, Any]]:
    """
    Extract concepts for a chapter. If content is very long, chunk and merge.
    """
    if not (content or "").strip():
        return []
    max_chunk = _max_chunk_tokens()
    toks = estimate_tokens(content)
    results: List[Dict[str, Any]] = []

    def _extract_on_text(txt: str, limit: int) -> List[Dict[str, Any]]:
        msgs = _build_concepts_messages(title, txt, limit)
        use_json = want_json_mode()
        if use_json:
            try:
                resp = client.chat.completions.create(
                    model=model,
                    messages=msgs,
                    temperature=0.2,
                    response_format={"type": "json_object"},
                )
                return _parse_concepts_json(resp.choices[0].message.content or "", limit)
            except Exception:
                pass
        resp = client.chat.completions.create(model=model, messages=msgs, temperature=0.2)
        return _parse_concepts_json(resp.choices[0].message.content or "", limit)

    if toks > max_chunk:
        chunks = split_text_by_tokens(content, max_chunk)
        # Allocate concept budget per chunk proportional to size
        chunk_tokens = [estimate_tokens(c) for c in chunks]
        total = sum(chunk_tokens) or 1
        alloc = [max(1, round(concepts_max * ct / total)) for ct in chunk_tokens]
        # Adjust to sum exactly concepts_max
        diff = sum(alloc) - concepts_max
        i = 0
        while diff != 0 and alloc:
            if diff > 0 and alloc[i] > 1:
                alloc[i] -= 1
                diff -= 1
            elif diff < 0:
                alloc[i] += 1
                diff += 1
            i = (i + 1) % len(alloc)
        for chunk, k in zip(chunks, alloc):
            part = _extract_on_text(chunk, k)
            results.extend(part)
    else:
        results = _extract_on_text(content, concepts_max)

    # Deduplicate by normalized title
    seen = set()
    merged: List[Dict[str, Any]] = []
    for c in results:
        t = (c.get("title") or "").strip().lower()
        if not t or t in seen:
            continue
        seen.add(t)
        merged.append(c)
    return merged[:concepts_max]


def _parse_qa_with_covers_json(text: str, max_items: int) -> List[Dict[str, Any]]:
    """
    Expect JSON: { "questions": [ { "question": "...", "answer": "...", "covers": ["C1","C2"] } ] }
    """
    try:
        data = json.loads(text)
    except Exception:
        start = text.find("{")
        end = text.rfind("}")
        if start != -1 and end != -1 and end > start:
            try:
                data = json.loads(text[start:end + 1])
            except Exception:
                return []
        else:
            return []
    arr = data.get("questions")
    if not isinstance(arr, list):
        return []
    out: List[Dict[str, Any]] = []
    for it in arr[:max_items]:
        if not isinstance(it, dict):
            continue
        q = (it.get("question") or "").strip()
        a = (it.get("answer") or "").strip()
        covers = it.get("covers")
        cov = [c.strip() for c in covers if isinstance(c, str)] if isinstance(covers, list) else []
        if q and a:
            out.append({"question": q, "answer": a, "covers": cov})
    return out


def _build_coverage_messages(title: str, content: str, concept_snippets: List[Dict[str, str]], want_n: int) -> List[Dict[str, str]]:
    """
    Build messages instructing the model to generate Q&A with covers IDs.
    concept_snippets: list of { "id": "C1", "title": "...", "summary": "..." }
    """
    system = (
        "You create question-answer pairs strictly grounded in the given chapter text. "
        "Each Q&A must clearly cover one or more of the provided concept IDs. Output only JSON."
    )
    concepts_text = "\n".join([f"- {snip['id']}: {snip['title']} — {snip.get('summary','')}" for snip in concept_snippets])
    user = (
        f"Chapter Title: {title}\n"
        f"Chapter Content:\n{content}\n\n"
        f"Uncovered concepts (use their IDs in 'covers'):\n{concepts_text}\n\n"
        f"Instructions:\n"
        f"- Generate up to {want_n} Q&A pairs covering the uncovered concepts (prefer high-importance if applicable).\n"
        "- JSON only, schema:\n"
        '{ "questions": [ { "question": "string", "answer": "string", "covers": ["C1","C2"] } ] }\n'
        "- No commentary or extra keys."
    )
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]


def generate_qa_for_chapter_coverage(
    client,
    model: str,
    title: str,
    content: str,
    coverage_threshold: float,
    concepts_max: int,
    max_qa_safety: int,
    max_iterations: int,
) -> List[Dict[str, str]]:
    """
    Coverage-driven Q&A generation. Returns final list of {"question","answer"}.
    """
    concepts = extract_concepts_for_chapter(client, model, title, content, concepts_max=concepts_max)
    if not concepts:
        return []

    # Assign IDs
    with_ids: List[Dict[str, Any]] = []
    for i, c in enumerate(concepts, start=1):
        with_ids.append({
            "id": f"C{i}",
            "title": c.get("title", ""),
            "summary": c.get("summary", ""),
            "importance": int(c.get("importance", 2)),
        })

    total = len(with_ids)
    uncovered_ids = {c["id"] for c in with_ids}
    qa_acc: List[Dict[str, str]] = []
    iterations = 0

    def _select_uncovered_slices(limit: int = 20) -> List[Dict[str, str]]:
        # Sort by importance desc, then by title
        ordered = sorted(
            [c for c in with_ids if c["id"] in uncovered_ids],
            key=lambda x: (x.get("importance", 2), x.get("title","")),
            reverse=True,
        )
        selected = ordered[:limit]
        return [{"id": c["id"], "title": c["title"], "summary": c.get("summary","")} for c in selected]

    while iterations < max_iterations and uncovered_ids and len(qa_acc) < max_qa_safety:
        iterations += 1
        # Decide how many to ask in this iteration
        remaining_concepts = len(uncovered_ids)
        want_n = min( max(3, round(remaining_concepts / 3)), max_qa_safety - len(qa_acc) )

        snippets = _select_uncovered_slices(limit=20)
        msgs = _build_coverage_messages(title, content, snippets, want_n)

        use_json = want_json_mode()
        if use_json:
            try:
                resp = client.chat.completions.create(
                    model=model,
                    messages=msgs,
                    temperature=0.3,
                    response_format={"type": "json_object"},
                )
                batch = _parse_qa_with_covers_json(resp.choices[0].message.content or "", max_items=want_n)
            except Exception:
                resp = client.chat.completions.create(model=model, messages=msgs, temperature=0.3)
                batch = _parse_qa_with_covers_json(resp.choices[0].message.content or "", max_items=want_n)
        else:
            resp = client.chat.completions.create(model=model, messages=msgs, temperature=0.3)
            batch = _parse_qa_with_covers_json(resp.choices[0].message.content or "", max_items=want_n)

        # Deduplicate by question text, and update coverage
        seen_q = { (q.get("question") or "").strip().lower() for q in qa_acc }
        new_items: List[Dict[str, str]] = []
        newly_covered: set = set()
        for it in batch or []:
            q = (it.get("question") or "").strip()
            a = (it.get("answer") or "").strip()
            if not q or not a:
                continue
            key = q.lower()
            if key in seen_q:
                continue
            new_items.append({"question": q, "answer": a})
            seen_q.add(key)
            for cid in (it.get("covers") or []):
                if isinstance(cid, str):
                    c = cid.strip()
                    if c:
                        newly_covered.add(c)

        qa_acc.extend(new_items)
        uncovered_ids -= newly_covered

        covered = (total - len(uncovered_ids))
        coverage = covered / total if total else 1.0
        print(f"[qa][coverage] '{title or 'Untitled'}': iter={iterations} added={len(new_items)} covered={covered}/{total} ({coverage:.0%})", flush=True)

        if coverage >= coverage_threshold:
            break

        # Stop if marginal progress stalls
        if len(new_items) == 0:
            break

    return qa_acc[:max_qa_safety]


# ----------------------------
# Chapter extraction
# ----------------------------

def _is_heading_style(style_name: str, target_level: int) -> bool:
    """Heuristic detection of heading styles including localized variants."""
    s = (style_name or "").strip().lower()
    # Common English: "heading 1"
    if f"heading {target_level}" in s:
        return True
    # Very simple heuristic for localized patterns (e.g., German "überschrift 1")
    # This won't cover all locales, but is better than nothing.
    if any(key in s for key in ["überschrift", "titre", "título", "título", "rubrique", "título", "titolo", "intestazione"]):
        if str(target_level) in s:
            return True
    # Exact match fallback
    return s == f"heading{target_level}"


def _get_outline_level(p) -> Optional[int]:
    """
    Try to read Word outline level (0-based) from paragraph XML.
    Returns int or None. Heading 1 -> 0, Heading 2 -> 1, etc.
    """
    try:
        ppr = p._p.pPr
        if ppr is not None and ppr.outlineLvl is not None:
            return int(ppr.outlineLvl.val)
    except Exception:
        return None
    return None


def split_into_chapters_python_docx(docx_path: str, heading_level: int = 1) -> List[Dict[str, Any]]:
    if Document is None:
        raise RuntimeError("python-docx is not installed. Please install 'python-docx' to use the fallback parser.")

    doc = Document(docx_path)
    chapters: List[Dict[str, Any]] = []

    current_title: Optional[str] = None
    current_text_parts: List[str] = []

    for p in doc.paragraphs:
        style_name = ""
        try:
            style_name = p.style.name or ""
        except Exception:
            style_name = ""

        text = (p.text or "").strip()

        is_head = (_is_heading_style(style_name, heading_level) or (_get_outline_level(p) == heading_level - 1))
        if is_head and text:
            # Close previous chapter
            if current_title is not None:
                content = "\n\n".join([t for t in current_text_parts if t.strip()])
                chapters.append({
                    "title": current_title,
                    "content": content
                })
            # Start new
            current_title = text
            current_text_parts = []
        else:
            # Accumulate body text
            if text:
                current_text_parts.append(text)

    # Flush last
    if current_title is None:
        # No headings found; treat entire document as one chapter
        full_text = "\n\n".join([(p.text or "").strip() for p in doc.paragraphs if (p.text or "").strip()])
        chapters.append({
            "title": "Document",
            "content": full_text
        })
    else:
        content = "\n\n".join([t for t in current_text_parts if t.strip()])
        chapters.append({
            "title": current_title,
            "content": content
        })

    # Clean up whitespace
    normalized = []
    for ch in chapters:
        title = (ch.get("title") or "").strip()
        content = re.sub(r"\n{3,}", "\n\n", (ch.get("content") or "").strip())
        normalized.append({"title": title, "content": content})
    return normalized


def split_into_chapters_docling(docx_path: str, heading_level: int = 1) -> Optional[List[Dict[str, Any]]]:
    """
    Placeholder docling integration. If docling is installed and usable, implement here.
    For now, return None to force fallback to python-docx unless you add the actual parser.
    """
    try:
        import docling  # type: ignore  # noqa: F401
        # NOTE: Implement real docling extraction if you prefer. Example outline:
        # converter = docling.DocumentConverter()
        # result = converter.convert(docx_path)
        # ... build chapters by heading hierarchy ...
        # Return list[{"title": "...", "content": "..."}]
        # For conservative reliability, we return None to use python-docx.
        return None
    except Exception:
        return None


def extract_chapters_for_file(docx_path: str, heading_level: int = 1) -> List[Dict[str, Any]]:
    # Try docling first (if implemented), else fallback to python-docx
    chapters = split_into_chapters_docling(docx_path, heading_level=heading_level)
    if not chapters:
        chapters = split_into_chapters_python_docx(docx_path, heading_level=heading_level)

    # If only one chapter detected, try auto-detecting a better heading level (1..3)
    try:
        if len(chapters) <= 1:
            best = chapters
            for lvl in (1, 2, 3):
                ch_try = split_into_chapters_python_docx(docx_path, heading_level=lvl)
                if len(ch_try) > len(best):
                    best = ch_try
            chapters = best
    except Exception:
        pass

    # Enrich with counts
    enriched: List[Dict[str, Any]] = []
    for idx, ch in enumerate(chapters, start=1):
        content = ch.get("content", "") or ""
        enriched.append({
            "index": idx,
            "title": ch.get("title", f"Chapter {idx}") or f"Chapter {idx}",
            "content": content,
            "char_count": len(content),
            "token_estimate": estimate_tokens(content)
        })
    return enriched


# ----------------------------
# LLM integration (Scaleway OpenAI-compatible)
# ----------------------------

def get_openai_client():
    global OpenAI
    if OpenAI is None:
        try:
            from openai import OpenAI as _OpenAI  # type: ignore
        except Exception as e:
            raise RuntimeError("The 'openai' package is required for the 'qa' step. Install it first.") from e
        OpenAI = _OpenAI

    # Resolve base URL from multiple env vars; normalize to ".../v1"
    raw_base = (os.getenv("SCW_BASE_URL") or os.getenv("SCW_OPENAI_BASE_URL") or os.getenv("OPENAI_BASE_URL") or os.getenv("OPENAI_API_BASE") or "").strip()
    api_key = (os.getenv("SCW_API_KEY") or os.getenv("SCW_SECRET_KEY") or os.getenv("OPENAI_API_KEY") or "").strip()

    if not raw_base:
        # Sensible default for Scaleway
        raw_base = "https://api.scaleway.ai"

    # Strip known chat path suffixes and ensure /v1
    base = re.sub(r"/(?:openai/)?v\d+/chat/completions/?$", "", raw_base.rstrip("/"), flags=re.IGNORECASE)
    base = re.sub(r"/providers/openai/chat/completions/?$", "", base.rstrip("/"), flags=re.IGNORECASE)
    base = re.sub(r"/chat/completions/?$", "", base.rstrip("/"), flags=re.IGNORECASE)
    if not re.search(r"/v\d+$", base):
        base = base.rstrip("/") + "/v1"

    if not api_key:
        raise RuntimeError("Missing API key. Set SCW_API_KEY (or SCW_SECRET_KEY / OPENAI_API_KEY).")

    client = OpenAI(base_url=base, api_key=api_key)
    return client


def get_model_name() -> str:
    model = os.getenv("SCW_MODEL") or os.getenv("LLM_MODEL") or os.getenv("TOGETHER_MODEL") or "gpt-oss-120b"
    model = model.strip()
    if "/" in model:
        parts = model.split("/")
        if parts[-1]:
            model = parts[-1]
    return model


def want_json_mode() -> bool:
    return os.getenv("SCW_JSON_MODE", "").strip() in ("1", "true", "True", "yes")


def build_messages_for_chapter(title: str, content: str, n: int) -> List[Dict[str, str]]:
    system = (
        "You are an expert knowledge designer. " \
        "Read the provided chapter text carefully and understand its content fully"
        "Based only on the provided chapter, "
        "generate highly relevant question and answer pairs that help readers understand the material. "
        "Try to cover key concepts, definitions, and insights from the chapter. "
        "Also try to generate questions, a user might realistically ask if he did NOT read the chapter. "
        "Ideally, generate questions in the 'how do I ...' format."
        "Our aim is to help users understand the chapter content through these Q&A pairs. "
        "and to mimic a realistic user curiosity. "
        "Answers must be grounded strictly in the chapter content."
    )
    user = (
        f"Chapter Title: {title}\n"
        f"Chapter Content:\n{content}\n\n"
        f"Instructions:\n"
        f"- Generate up to {n} unique Q&A pairs.\n"
        f"- Each question must be answerable strictly from the chapter content.\n"
        f"- Keep each answer concise (2–4 sentences).\n"
        f"- Output JSON with this exact schema:\n"
        f"{{\n"
        f'  "questions": [\n'
        f'    {{ "question": "string", "answer": "string" }}\n'
        f"  ]\n"
        f"}}\n"
        f"No additional keys, no markdown, no comments."
    )
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]


def _parse_questions_json(text: str, max_items: int) -> List[Dict[str, str]]:
    # Attempt to parse a JSON object with "questions" list
    def try_load(s: str) -> Optional[Dict[str, Any]]:
        try:
            return json.loads(s)
        except Exception:
            return None

    data = try_load(text)
    if data is None:
        # Extract the largest JSON object in the text
        start = text.find("{")
        end = text.rfind("}")
        if start != -1 and end != -1 and end > start:
            data = try_load(text[start:end + 1])

    if not isinstance(data, dict):
        return []

    questions = data.get("questions")
    if not isinstance(questions, list):
        return []

    cleaned: List[Dict[str, str]] = []
    for item in questions[:max_items]:
        q = (item.get("question") if isinstance(item, dict) else None) if item is not None else None
        a = (item.get("answer") if isinstance(item, dict) else None) if item is not None else None
        if isinstance(q, str) and isinstance(a, str):
            q = q.strip()
            a = a.strip()
            if q and a:
                cleaned.append({"question": q, "answer": a})
    return cleaned


@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=1, max=10))
def generate_qa_for_chapter(client, model: str, title: str, content: str, n: int) -> List[Dict[str, str]]:
    # Safeguard against empty content
    if not (content or "").strip():
        return []

    messages = build_messages_for_chapter(title, content, n)
    use_json = want_json_mode()

    # Try JSON mode first if requested
    if use_json:
        try:
            resp = client.chat.completions.create(
                model=model,
                messages=messages,
                temperature=0.3,
                response_format={"type": "json_object"},
            )
            text = resp.choices[0].message.content or ""
            items = _parse_questions_json(text, max_items=n)
            if items:
                return items
        except Exception:
            # Fall back to non-JSON mode below
            pass

    # Non-JSON mode
    resp = client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=0.3,
    )
    text = resp.choices[0].message.content or ""
    items = _parse_questions_json(text, max_items=n)
    return items


# ----------------------------
# Commands
# ----------------------------

def cmd_extract(args: argparse.Namespace) -> int:
    input_dir = args.input
    output_dir = args.output
    heading_level = args.heading_level

    if not os.path.isdir(input_dir):
        print(f"[extract] Input directory does not exist: {input_dir}", file=sys.stderr)
        return 2

    ensure_dir(output_dir)
    files = list_docx_files(input_dir)
    if not files:
        print(f"[extract] No .docx files found in {input_dir}", file=sys.stderr)
        return 1

    for path in files:
        base = os.path.splitext(os.path.basename(path))[0]
        out_path = os.path.join(output_dir, f"{base}.json")

        chapters = extract_chapters_for_file(path, heading_level=heading_level)
        out = {
            "source_file": os.path.basename(path),
            "heading_level": heading_level,
            "extracted_at": iso_now(),
            "chapters": chapters,
        }

        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(out, f, ensure_ascii=False, indent=2)

        print(f"[extract] Wrote chapters JSON: {out_path} (chapters={len(chapters)})")

    return 0


def cmd_qa(args: argparse.Namespace) -> int:
    input_dir = args.input
    output_dir = args.output
    max_per_chapter = args.max_per_chapter

    if not os.path.isdir(input_dir):
        print(f"[qa] Input directory does not exist: {input_dir}", file=sys.stderr)
        return 2

    ensure_dir(output_dir)

    # Lazy-init client so extract can run without openai installed
    client = get_openai_client()
    model = get_model_name()

    # Read all chapters JSON files
    json_files = [os.path.join(input_dir, name) for name in os.listdir(input_dir) if name.lower().endswith(".json")]
    json_files.sort()
    if not json_files:
        print(f"[qa] No chapters JSON files found in {input_dir}", file=sys.stderr)
        return 1

    for chapter_json_path in json_files:
        with open(chapter_json_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        source_file = data.get("source_file") or "unknown.docx"
        base = os.path.splitext(os.path.basename(source_file))[0]
        out_path = os.path.join(output_dir, f"QA-{base}.json")

        chapters = data.get("chapters") or []
        print(f"[qa] Processing {base}: {len(chapters)} chapters", flush=True)

        # Optional per-document cap allocation
        max_per_document = getattr(args, "max_per_document", None)
        allocations: Optional[List[int]] = None
        if max_per_document is not None:
            # Compute base demand and final allocations for logging
            base_demands: List[int] = []
            for ch in chapters:
                te_dbg = int(ch.get("token_estimate") or estimate_tokens((ch.get("content") or "")))
                approx_dbg = round(te_dbg / 350.0)
                base_demands.append(min(max_per_chapter, max(0, approx_dbg)))
            allocations = compute_chapter_allocations(chapters, max_per_chapter, max_per_document)
            print(f"[qa]  doc-cap: requested={sum(base_demands)}, cap={max_per_document}, assigned={sum(allocations)}", flush=True)

        results: List[Dict[str, Any]] = []
        for idx_ch, ch in enumerate(chapters, start=1):
            title = (ch.get("title") or "").strip()
            content = (ch.get("content") or "").strip()
            token_est = int(ch.get("token_estimate") or estimate_tokens(content))

            # Coverage-mode branch
            if getattr(args, "coverage_mode", False):
                print(f"[qa]  - Chapter {idx_ch}/{len(chapters)}: '{title or 'Untitled'}' tokens≈{token_est} [coverage-mode]", flush=True)
                try:
                    questions = generate_qa_for_chapter_coverage(
                        client=client,
                        model=model,
                        title=title or "Chapter",
                        content=content,
                        coverage_threshold=float(getattr(args, "coverage_threshold", 0.85)),
                        concepts_max=int(getattr(args, "concepts_max", 50)),
                        max_qa_safety=int(getattr(args, "max_qa_per_chapter_safety", 60)),
                        max_iterations=int(getattr(args, "max_iterations", 8)),
                    )
                except Exception as e:
                    print(f"[qa] LLM error (coverage) for chapter '{title}': {e}", file=sys.stderr)
                    questions = []
            else:
                # Existing size-based target
                if 'allocations' in locals() and allocations is not None:
                    n = allocations[idx_ch - 1]
                else:
                    n = determine_qa_count(token_est, max_per_chapter=max_per_chapter)
                tag = " (doc-cap)" if 'allocations' in locals() and allocations is not None else ""
                print(f"[qa]  - Chapter {idx_ch}/{len(chapters)}: '{title or 'Untitled'}' tokens≈{token_est}, target={n}{tag}", flush=True)

                if not content:
                    print(f"[qa]    -> skipped (empty content)", flush=True)
                    results.append({
                        "chapter_title": title or "Untitled",
                        "question_count": 0,
                        "questions": []
                    })
                    continue
                if 'allocations' in locals() and allocations is not None and n <= 0:
                    print(f"[qa]    -> allocated 0 (doc-cap), skipping", flush=True)
                    results.append({
                        "chapter_title": title or "Untitled",
                        "question_count": 0,
                        "questions": []
                    })
                    continue

                try:
                    max_chunk = _max_chunk_tokens()
                    total_tokens = estimate_tokens(content)
                    if total_tokens > max_chunk:
                        chunks = split_text_by_tokens(content, max_chunk)
                        print(f"[qa]    -> chunking into {len(chunks)} parts for context", flush=True)
                        # Distribute n across chunks proportionally to chunk size
                        chunk_tokens = [estimate_tokens(c) for c in chunks]
                        sum_tokens = sum(chunk_tokens) or 1
                        alloc = [max(1, round(n * ct / sum_tokens)) for ct in chunk_tokens]
                        # Adjust allocation to sum exactly n
                        diff = sum(alloc) - n
                        i_adj = 0
                        while diff != 0 and alloc:
                            if diff > 0 and alloc[i_adj] > 0:
                                alloc[i_adj] -= 1
                                diff -= 1
                            elif diff < 0:
                                alloc[i_adj] += 1
                                diff += 1
                            i_adj = (i_adj + 1) % len(alloc)
                        all_q: List[Dict[str, str]] = []
                        for i_chunk, (chunk, k) in enumerate(zip(chunks, alloc), start=1):
                            if k <= 0:
                                continue
                            sub = generate_qa_for_chapter(client, model, title or "Chapter", chunk, k)
                            all_q.extend(sub or [])
                        # Dedupe by question text and clamp
                        seen_q = set()
                        deduped: List[Dict[str, str]] = []
                        for qa in all_q:
                            qtxt = (qa.get("question") or "").strip()
                            atxt = (qa.get("answer") or "").strip()
                            if qtxt and atxt:
                                key = qtxt.lower()
                                if key not in seen_q:
                                    seen_q.add(key)
                                    deduped.append({"question": qtxt, "answer": atxt})
                        questions = deduped[:n]
                    else:
                        questions = generate_qa_for_chapter(client, model, title or "Chapter", content, n)
                except Exception as e:
                    print(f"[qa] LLM error for chapter '{title}': {e}", file=sys.stderr)
                    questions = []

            # Normalize and (optionally) clamp
            questions = [{"question": q.get("question", "").strip(), "answer": q.get("answer", "").strip()}
                         for q in (questions or []) if q.get("question") and q.get("answer")]
            if not getattr(args, "coverage_mode", False):
                target_n = (allocations[idx_ch - 1] if 'allocations' in locals() and allocations is not None else determine_qa_count(token_est, max_per_chapter=max_per_chapter))
                if len(questions) > target_n:
                    questions = questions[:target_n]
            print(f"[qa]    -> got {len(questions)} items", flush=True)

            results.append({
                "chapter_title": title or "Untitled",
                "question_count": len(questions),
                "questions": questions
            })

        final_out = {
            "source_file": source_file,
            "model": model,
            "generated_at": iso_now(),
            "chapters": results
        }

        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(final_out, f, ensure_ascii=False, indent=2)

        print(f"[qa] Wrote Q&A JSON: {out_path}")
        print(f"[qa] Completed {base}: total rows={sum(len(r.get('questions', [])) for r in results)}", flush=True)

    return 0


def cmd_docx(args: argparse.Namespace) -> int:
    input_dir = args.input
    output_dir = args.output

    if Document is None:
        raise RuntimeError("python-docx is required to write DOCX files. Please install 'python-docx'.")

    if not os.path.isdir(input_dir):
        print(f"[docx] Input directory does not exist: {input_dir}", file=sys.stderr)
        return 2

    ensure_dir(output_dir)

    json_files = [os.path.join(input_dir, name) for name in os.listdir(input_dir) if name.lower().endswith(".json")]
    json_files.sort()
    if not json_files:
        print(f"[docx] No Q&A JSON files found in {input_dir}", file=sys.stderr)
        return 1

    for json_path in json_files:
        try:
            with open(json_path, "r", encoding="utf-8") as f:
                data = json.load(f)
        except Exception as e:
            print(f"[docx] Failed to read {json_path}: {e}", file=sys.stderr)
            continue

        source_file = data.get("source_file") or os.path.basename(json_path)
        base = os.path.splitext(os.path.basename(source_file))[0]
        out_path = os.path.join(output_dir, f"QA-{base}.docx")

        doc = Document()
        doc.add_heading(f"Q&A for {base}", level=1)

        model = data.get("model") or ""
        generated_at = data.get("generated_at") or ""
        meta = " | ".join([p for p in [f"Model: {model}" if model else "", f"Generated: {generated_at}" if generated_at else ""] if p])
        if meta:
            doc.add_paragraph(meta)

        table = doc.add_table(rows=1, cols=2)
        hdr = table.rows[0].cells
        hdr[0].text = "Question"
        hdr[1].text = "Answer"

        rows_added = 0
        for ch in data.get("chapters", []):
            chapter_title = (ch.get("chapter_title") or ch.get("title") or "Untitled").strip()
            for qa in ch.get("questions", []):
                q = (qa.get("question") or "").strip()
                a = (qa.get("answer") or "").strip()
                if not q or not a:
                    continue
                row = table.add_row().cells
                row[0].text = q
                row[1].text = a
                rows_added += 1

        if rows_added == 0:
            doc.add_paragraph("No questions found for this document.")

        try:
            doc.save(out_path)
            print(f"[docx] Wrote DOCX: {out_path} (rows={rows_added})")
        except Exception as e:
            print(f"[docx] Failed to write {out_path}: {e}", file=sys.stderr)

    return 0


# ----------------------------
# Main
# ----------------------------

def build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(description="Generate Q&A JSON from DOCX via a two-step pipeline.")
    sub = p.add_subparsers(dest="command", required=True)

    # extract
    pe = sub.add_parser("extract", help="Extract chapters from DOCX into intermediate JSON files.")
    pe.add_argument("--input", required=True, help="Input directory containing .docx files (e.g., QA_SOURCE)")
    pe.add_argument("--output", required=True, help="Output directory for chapters JSON (e.g., QA_CHAPTERS)")
    pe.add_argument("--heading-level", type=int, default=1, help="Heading level to treat as 'chapter' (default: 1)")
    pe.set_defaults(func=cmd_extract)

    # qa
    pq = sub.add_parser("qa", help="Generate Q&A from chapters JSON using an OpenAI-compatible endpoint.")
    pq.add_argument("--input", required=True, help="Input directory containing chapters JSON (e.g., QA_CHAPTERS)")
    pq.add_argument("--output", required=True, help="Output directory for final Q&A JSON (e.g., QA_OUTPUT)")
    pq.add_argument("--max-per-chapter", type=int, default=10, help="Maximum Q&A items per chapter (default: 10)")
    pq.add_argument("--max-per-document", type=int, default=None, help="Optional cap on total Q&A pairs per document (distributed by chapter size)")
    # Coverage-mode options
    pq.add_argument("--coverage-mode", action="store_true", help="Use concept coverage-driven generation (adaptive count per chapter)")
    pq.add_argument("--coverage-threshold", type=float, default=0.85, help="Target concept coverage ratio (0..1), default 0.85")
    pq.add_argument("--concepts-max", type=int, default=50, help="Max concepts to extract per chapter (default 50)")
    pq.add_argument("--max-qa-per-chapter-safety", type=int, default=60, help="Safety cap per chapter in coverage mode (default 60)")
    pq.add_argument("--max-iterations", type=int, default=8, help="Max coverage iterations per chapter (default 8)")
    pq.set_defaults(func=cmd_qa)

    # docx
    pd = sub.add_parser("docx", help="Create a DOCX per Q&A JSON with a table of chapter, question, answer.")
    pd.add_argument("--input", required=True, help="Input directory containing final Q&A JSON (e.g., QA_OUTPUT)")
    pd.add_argument("--output", required=True, help="Output directory for DOCX files (e.g., QA_DOCX)")
    pd.set_defaults(func=cmd_docx)

    return p


def main(argv: Optional[List[str]] = None) -> int:
    # Load environment from .env if present
    load_env()
    parser = build_parser()
    args = parser.parse_args(argv)
    try:
        return args.func(args)
    except RuntimeError as e:
        print(str(e), file=sys.stderr)
        return 2
    except KeyboardInterrupt:
        print("Interrupted.", file=sys.stderr)
        return 130


if __name__ == "__main__":
    sys.exit(main())
