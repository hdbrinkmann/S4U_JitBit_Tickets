#!/usr/bin/env python3
"""
Generate DOCX files from Ticket_Data.JSON with configurable tickets per file.

What this script does:
- Reads a top-level JSON array of tickets with keys:
  ticket_id, date, subject, problem, solution, image_urls
- Groups tickets into batches (default: 50 tickets per DOCX file)
- Renders each ticket with one page per ticket in the DOCX file:
  - Subject as heading
  - "Problem" section
  - "Lösung" section
  - Images listed in image_urls embedded below "Problem" (optional)
- Supports image URLs from multiple sources:
  * JitBit-protected images: uses API-first fetching with Bearer token:
    GET {base}/helpdesk/api/attachment?id={FileID}
    falling back to {base}/api/attachment?id=...
    where {base} is derived from JITBIT_BASE_URL or provided args
  * Jira attachment URLs: directly downloads from URLs like:
    https://{instance}.atlassian.net/rest/api/3/attachment/content/{id}
  * Generic external images: standard HTTP(S) fetch

Usage:
  1) Ensure .env contains appropriate tokens:
       JITBIT_API_TOKEN=...     (Bearer token for Jitbit API)
       JITBIT_BASE_URL=https://support.example.com/helpdesk
       JIRA_API_TOKEN=...       (Bearer token for Jira API, optional)
     Note: Ticket_Data.JSON usually doesn't contain export_info.api_base_url,
           so we rely on JITBIT_BASE_URL to derive the API root.

  2) Run:
     # Default: 50 tickets per DOCX file
     python tickets_to_docx.py --input Ticket_Data.JSON --verbose true
     
     # Custom: 25 tickets per DOCX file
     python tickets_to_docx.py --input Ticket_Data.JSON --tickets-per-file 25
     
     # One ticket per file (original behavior)
     python tickets_to_docx.py --input Ticket_Data.JSON --tickets-per-file 1
"""

import argparse
import io
import json
import os
import re
import sys
import base64
from typing import List, Optional, Tuple

try:
    from dotenv import load_dotenv
    load_dotenv()
except Exception:
    pass

from PIL import Image as PILImage

# python-docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Lightweight helpers and fetcher (decoupled from ReportLab deps)
import requests
from urllib.parse import urljoin, urlparse, parse_qs
from concurrent.futures import ThreadPoolExecutor, as_completed
import hashlib
import warnings
import math
import time
from datetime import datetime, timedelta
warnings.filterwarnings("ignore", message="Palette images with Transparency", category=UserWarning)
warnings.filterwarnings("ignore", category=UserWarning, module="PIL.Image")

def xml_safe(s: Optional[str]) -> str:
    """
    Remove XML 1.0 illegal control characters (except TAB, LF, CR) that cause python-docx to fail.
    """
    if not s:
        return ""
    return re.sub(r"[\x00-\x08\x0B-\x0C\x0E-\x1F]", "", s)

def str2bool(v: str) -> bool:
    return str(v).lower() in {"1", "true", "t", "yes", "y"}

def sanitize_url(raw: Optional[str]) -> Optional[str]:
    if raw is None:
        return None
    s = str(raw).strip()
    if not s:
        return None
    # Strip surrounding quotes
    if (s.startswith('"') and s.endswith('"')) or (s.startswith("'") and s.endswith("'")):
        s = s[1:-1]
    # Cut at first whitespace
    s = s.split()[0]
    # Remove any embedded tags
    for ch in ("<", ">"):
        pos = s.find(ch)
        if pos != -1:
            s = s[:pos]
    return s or None

def resolve_url(src: str, base: Optional[str]) -> Optional[str]:
    if not src:
        return None
    src = src.strip()
    if src.startswith("http://") or src.startswith("https://"):
        return src
    if base:
        return urljoin(base.rstrip("/") + "/", src.lstrip("/"))
    return None

def extract_file_id_from_url(url_str: str) -> Optional[str]:
    """
    Extracts a numeric FileID from typical Jitbit URLs:
      - /helpdesk/File/Get/26355
      - /helpdesk/File/Get?id=26355
    Returns None if no numeric id present.
    """
    try:
        pu = urlparse(url_str)
        # Query ?id=123
        qs = parse_qs(pu.query or "")
        if "id" in qs:
            v = qs["id"][0]
            if str(v).isdigit():
                return str(v)
        # Path segment
        segs = [s for s in (pu.path or "").split("/") if s]
        if segs:
            last = segs[-1]
            if last.isdigit():
                return last
    except Exception:
        pass
    return None

def is_jira_attachment_url(url_str: str) -> bool:
    """
    Detects if a URL is a Jira attachment URL.
    Handles both REST API download URLs and /secure/attachment paths, e.g.:
      - https://{instance}.atlassian.net/rest/api/3/attachment/{id}/content
      - https://{instance}.atlassian.net/rest/api/latest/attachment/content/{id}
      - https://{instance}.atlassian.net/secure/attachment/{id}/{filename}
      - https://{instance}.atlassian.net/secure/thumbnail/{id}/{filename}
    """
    try:
        pu = urlparse(url_str)
        host_ok = bool(pu.hostname and ".atlassian.net" in pu.hostname)
        path = (pu.path or "").lower()
        api_pattern = ("/rest/api/" in path) and ("/attachment" in path) and ("/content" in path)
        secure_pattern = ("/secure/attachment/" in path) or ("/secure/thumbnail/" in path)
        return bool(host_ok and (api_pattern or secure_pattern))
    except Exception:
        return False

def derive_api_root(api_base_url: Optional[str], env_base_url: Optional[str], verbose: bool = False) -> Optional[str]:
    """
    Derive the API root:
      If api_base_url path contains '/helpdesk' -> {scheme}://{host}/helpdesk/api
      Else -> {scheme}://{host}/api
    Falls back to env_base_url if api_base_url is missing.
    """
    base = api_base_url or env_base_url
    if not base:
        return None
    try:
        pu = urlparse(base)
        scheme = pu.scheme or "https"
        netloc = pu.netloc
        path = (pu.path or "").lower()
        if not netloc:
            return None
        if "/helpdesk" in path:
            root = f"{scheme}://{netloc}/helpdesk/api"
        else:
            root = f"{scheme}://{netloc}/api"
        if verbose:
            print(f"[INFO] API root derived: {root}", file=sys.stderr)
        return root
    except Exception:
        return None

class AttachmentFetcher:
    def __init__(self, api_root: Optional[str], token: Optional[str], jira_token: Optional[str] = None, jira_email: Optional[str] = None, timeout: float = 15.0, verbose: bool = False):
        self.api_root = api_root
        # Jitbit API bearer token (if provided)
        self.token = token

        # Jira Cloud auth
        self.jira_api_token = jira_token
        self.jira_email = jira_email
        self.jira_basic_auth_header: Optional[str] = None
        if self.jira_email and self.jira_api_token:
            try:
                b64 = base64.b64encode(f"{self.jira_email}:{self.jira_api_token}".encode("utf-8")).decode("utf-8")
                self.jira_basic_auth_header = f"Basic {b64}"
            except Exception:
                self.jira_basic_auth_header = None

        self.timeout = timeout
        self.verbose = verbose

        # Session primarily for Jitbit fetches (Bearer token)
        self.session = requests.Session()
        self.session.headers.update({
            "User-Agent": "Tickets-DOCX/1.0",
            "Accept": "*/*",
        })
        if token:
            self.session.headers["Authorization"] = f"Bearer {token}"

    def _warn(self, msg: str):
        if self.verbose:
            print(msg, file=sys.stderr)

    def fetch_attachment_by_id(self, file_id: str) -> Optional[bytes]:
        """
        Fetches an attachment by FileID trying the correct API root(s).
        Many Jitbit installs live under '/helpdesk', so we try:
          1) {scheme}://{host}/helpdesk/api/attachment?id=...
          2) {scheme}://{host}/api/attachment?id=...
        We prioritize any '/helpdesk' root if self.api_root includes it.
        """
        if not self.api_root or not self.token:
            self._warn(f"[WARN] Missing API root or token for Jitbit attachment id={file_id}")
            return None

        try:
            p = urlparse(self.api_root)
            scheme = p.scheme or "https"
            netloc = p.netloc
            path = (p.path or "").lower()
            if not netloc:
                self._warn(f"[WARN] Invalid api_root (no host): {self.api_root}")
                return None

            candidates = []
            # Prioritize '/helpdesk/api' if api_root path mentions helpdesk
            if "/helpdesk" in path:
                candidates.append(f"{scheme}://{netloc}/helpdesk/api")
            # Always try '/helpdesk/api' first for safety if not already present
            if f"{scheme}://{netloc}/helpdesk/api" not in candidates:
                candidates.append(f"{scheme}://{netloc}/helpdesk/api")
            # Then plain '/api'
            candidates.append(f"{scheme}://{netloc}/api")

            last_err: Optional[Exception] = None
            for root in candidates:
                url = f"{root.rstrip('/')}/attachment?id={file_id}"
                try:
                    self._warn(f"[INFO] GET {url}")
                    r = self.session.get(url, timeout=self.timeout)
                    r.raise_for_status()
                    return r.content
                except Exception as e:
                    last_err = e
                    self._warn(f"[INFO] Candidate failed: {e}")
                    continue

            if last_err:
                self._warn(f"[WARN] API fetch failed for id={file_id}: {last_err}")
            return None
        except Exception as e:
            self._warn(f"[WARN] API fetch failed for id={file_id}: {e}")
            return None

    def fetch_jira_attachment(self, url: str) -> Optional[bytes]:
        """
        Fetch a Jira attachment or secure image URL using Jira Cloud authentication.
        Prefers Basic auth with JIRA_EMAIL + JIRA_API_TOKEN, falls back to Bearer if only token is provided.
        """
        try:
            # Separate session for Jira to isolate auth headers
            session = requests.Session()
            session.headers.update({
                "User-Agent": "Tickets-DOCX/1.0",
                "Accept": "*/*",
            })

            # Prefer Basic auth for Jira Cloud
            if self.jira_basic_auth_header:
                session.headers["Authorization"] = self.jira_basic_auth_header
                session.headers["X-Atlassian-Token"] = "no-check"
            elif self.jira_api_token:
                # Fallback to Bearer if only token was provided (may not work on Cloud)
                session.headers["Authorization"] = f"Bearer {self.jira_api_token}"

            self._warn(f"[INFO] GET (Jira) {url}")
            r = session.get(url, timeout=self.timeout, allow_redirects=True)
            r.raise_for_status()

            return r.content
        except Exception as e:
            self._warn(f"[WARN] Jira image fetch failed for {url}: {e}")
            return None

    def fetch_generic_image(self, url: str) -> Optional[bytes]:
        try:
            self._warn(f"[INFO] GET (external) {url}")
            r = self.session.get(url, timeout=self.timeout)
            r.raise_for_status()
            return r.content
        except Exception as e:
            self._warn(f"[WARN] External image fetch failed for {url}: {e}")
            return None


# ---- Page setup helpers ----

A4_INCH = (8.27, 11.69)
LETTER_INCH = (8.5, 11.0)
EMU_PER_INCH = 914400


def set_page_size_and_margins(doc: Document, page: str, margin_pt: float) -> None:
    """
    Configure the document's first section to the requested page size and margins.
    margin_pt is in points (1/72 inch).
    """
    section = doc.sections[0]

    if page.upper() == "A4":
        w_in, h_in = A4_INCH
    else:
        w_in, h_in = LETTER_INCH

    section.page_width = Inches(w_in)
    section.page_height = Inches(h_in)

    m_in = float(margin_pt) / 72.0
    section.left_margin = Inches(m_in)
    section.right_margin = Inches(m_in)
    section.top_margin = Inches(m_in)
    section.bottom_margin = Inches(m_in)


def get_usable_emu(doc: Document) -> Tuple[int, int]:
    """
    Return (usable_width_emu, usable_height_emu) for the first section.
    """
    s = doc.sections[0]
    usable_w = int(s.page_width - s.left_margin - s.right_margin)
    usable_h = int(s.page_height - s.top_margin - s.bottom_margin)
    return usable_w, usable_h


# ---- Text rendering helpers ----

_bullet_re = re.compile(r"^\s*[-\*\u2022]\s+(.*)")  # -, *, •
_ordered_re = re.compile(r"^\s*(\d+)[\.\)]\s+(.*)")  # 1. or 1)


def _iter_runs_from_markup(text: str):
    """
    Yields (segment, is_bold) by converting **bold** and <b>...</b> to runs.
    """
    if not text:
        return
    text = xml_safe(text)
    # normalize windows line breaks to avoid surprises
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Convert simple HTML bold to **bold** (handles escaped <b> too)
    text = re.sub(r"<b>(.*?)</b>", r"**\1**", text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r"<b>(.*?)</b>", r"**\1**", text, flags=re.IGNORECASE | re.DOTALL)

    pattern = re.compile(r"\*\*(.+?)\*\*")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            yield text[pos:m.start()], False
        yield m.group(1), True
        pos = m.end()
    if pos < len(text):
        yield text[pos:], False


def add_paragraph_with_inline_formatting(doc: Document, text: str, style: Optional[str] = None, color_rgb: Optional[RGBColor] = None):
    """
    Adds a paragraph to doc, splitting 'text' into bold/non-bold runs based on **...** or <b>...</b>.
    Returns the created paragraph.
    """
    p = doc.add_paragraph()
    if style:
        p.style = style
    for seg, is_bold in _iter_runs_from_markup(text or ""):
        run = p.add_run(seg)
        run.bold = bool(is_bold)
        if color_rgb is not None:
            run.font.color.rgb = color_rgb
    return p


def add_plain_text_block(doc: Document, text: str):
    """
    Convert plain text with paragraphs separated by blank lines.
    Supports simple unordered (-, *, •) and ordered (1., 2)) lists.
    Also supports **bold** inline.
    """
    if not (text and str(text).strip()):
        doc.add_paragraph("(No content)")
        return

    lines = str(text).replace("\r\n", "\n").replace("\r", "\n").split("\n")
    i = 0
    n = len(lines)

    while i < n:
        # skip blank lines
        while i < n and not lines[i].strip():
            i += 1
        if i >= n:
            break

        # detect list block
        m_b = _bullet_re.match(lines[i])
        m_o = _ordered_re.match(lines[i])
        if m_b or m_o:
            is_ordered = bool(m_o)
            style = "List Number" if is_ordered else "List Bullet"
            while i < n:
                mb = _bullet_re.match(lines[i])
                mo = _ordered_re.match(lines[i])
                if is_ordered and not mo:
                    break
                if (not is_ordered) and not mb:
                    break
                content = (mo.group(2) if mo else mb.group(1)).strip()
                p = add_paragraph_with_inline_formatting(doc, content)
                p.style = style
                i += 1
            continue

        # otherwise accumulate a paragraph until blank or next list
        para_lines: List[str] = []
        while i < n:
            if not lines[i].strip():
                i += 1
                break
            if _bullet_re.match(lines[i]) or _ordered_re.match(lines[i]):
                break
            para_lines.append(lines[i].strip())
            i += 1

        paragraph_text = " ".join(para_lines).strip()
        if paragraph_text:
            add_paragraph_with_inline_formatting(doc, paragraph_text)


# ---- Image helpers ----

def _bytes_to_image_dims_emu(img_bytes: bytes) -> Optional[Tuple[int, int]]:
    """
    Returns (width_emu, height_emu) using image DPI metadata if available, else assumes 96 DPI.
    """
    try:
        with PILImage.open(io.BytesIO(img_bytes)) as im:
            # Normalize palette+transparency PNGs to RGBA to avoid PIL warning
            if (im.format or "").upper() == "PNG" and im.mode == "P" and (im.info.get("transparency", None) is not None):
                im = im.convert("RGBA")
            w_px, h_px = im.width, im.height
            dpi = im.info.get("dpi", (96, 96))
            dpi_x = float(dpi[0] or 96.0)
            dpi_y = float(dpi[1] or 96.0)
            w_in = w_px / dpi_x
            h_in = h_px / dpi_y
            return int(Emu(Inches(w_in))), int(Emu(Inches(h_in)))
    except Exception:
        return None


def optimize_image_bytes(
    data: bytes,
    max_width_px: Optional[int],
    jpeg_quality: int = 75,
    convert_png_to_jpeg: bool = True,
    keep_png_if_transparent: bool = True,
    force_jpeg: bool = False,
    force_recompress_min_bytes: int = 128 * 1024,
    jpeg_optimize: bool = False,
    jpeg_progressive: bool = False,
    png_compress_level: int = 6,
) -> bytes:
    """
    Downscale and recompress image bytes to reduce DOCX size.
    - Optionally convert PNG to JPEG when transparency is not required.
    - Strip metadata by re-saving.
    - Keep original bytes if optimization does not reduce size.
    Fast-paths are included to avoid unnecessary work on small images.
    """
    try:
        with PILImage.open(io.BytesIO(data)) as im:
            fmt = (im.format or "").upper()
            # Normalize palette transparency to RGBA to avoid PIL warning and preserve alpha
            if fmt == "PNG" and im.mode == "P" and (im.info.get("transparency", None) is not None):
                im = im.convert("RGBA")
            has_alpha = (im.mode in ("RGBA", "LA")) or ("transparency" in im.info)

            # Fast path: skip tiny PNGs if not exceeding width and below recompress threshold
            w_px, h_px = im.width, im.height
            if (not force_jpeg) and fmt == "PNG" and (not max_width_px or w_px <= max_width_px) and len(data) < force_recompress_min_bytes:
                return data

            # Resize if larger than allowed (use thumbnail for speed and low memory)
            target_w = max_width_px or w_px
            if target_w and w_px > target_w:
                im = im.copy()
                target_h = max(1, int(h_px * (target_w / float(w_px))))
                im.thumbnail((target_w, target_h), PILImage.LANCZOS)
                w_px, h_px = im.width, im.height

            out = io.BytesIO()
            optimized: bytes

            if fmt == "PNG" and convert_png_to_jpeg and not (keep_png_if_transparent and has_alpha):
                im2 = im.convert("RGB")
                im2.save(
                    out,
                    format="JPEG",
                    quality=jpeg_quality,
                    optimize=jpeg_optimize,
                    progressive=jpeg_progressive,
                    subsampling=2,
                )
                optimized = out.getvalue()
            elif fmt in ("JPEG", "JPG"):
                # Recompress only if reasonably large or resized
                if len(data) >= force_recompress_min_bytes or (max_width_px and w_px > max_width_px):
                    im2 = im.convert("RGB")
                    im2.save(
                        out,
                        format="JPEG",
                        quality=jpeg_quality,
                        optimize=jpeg_optimize,
                        progressive=jpeg_progressive,
                        subsampling=2,
                    )
                    optimized = out.getvalue()
                else:
                    return data
            else:
                # Other formats (e.g., GIF, BMP, WEBP). If forcing JPEG, convert; else re-encode as PNG and keep only if smaller.
                if force_jpeg:
                    im2 = im.convert("RGB")
                    im2.save(
                        out,
                        format="JPEG",
                        quality=jpeg_quality,
                        optimize=jpeg_optimize,
                        progressive=jpeg_progressive,
                        subsampling=2,
                    )
                else:
                    im.save(out, format="PNG", optimize=False, compress_level=int(max(0, min(9, png_compress_level))))
                optimized = out.getvalue()

            # Only keep optimized if it actually reduces size, unless forcing JPEG conversion
            return optimized if (force_jpeg or len(optimized) < len(data)) else data
    except Exception:
        return data


def average_hash(data: bytes, hash_size: int = 8) -> Optional[int]:
    try:
        with PILImage.open(io.BytesIO(data)) as im:
            # Normalize palette+transparency PNGs to RGBA before grayscale convert
            if (im.format or "").upper() == "PNG" and im.mode == "P" and (im.info.get("transparency", None) is not None):
                im = im.convert("RGBA")
            im = im.convert("L")
            im = im.resize((hash_size, hash_size), PILImage.BILINEAR)
            pixels = list(im.getdata())
            avg = sum(pixels) / len(pixels)
            bits = 0
            for i, p in enumerate(pixels):
                if p >= avg:
                    bits |= (1 << i)
            return bits
    except Exception:
        return None

def hamming_distance(a: int, b: int) -> int:
    try:
        return (a ^ b).bit_count()
    except Exception:
        x = a ^ b
        count = 0
        while x:
            x &= x - 1
            count += 1
        return count

def add_image_placeholder_docx(doc: Document, url: str, auth_hint: bool = False, label: Optional[str] = None):
    safe_url = xml_safe(str(url or "").strip())
    label_text = f"{label} – " if label else ""
    hint_text = " (login/cookies may be required)" if auth_hint else ""
    p = doc.add_paragraph()
    run = p.add_run(xml_safe(f"{label_text}Image could not be loaded: {safe_url}{hint_text}"))
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def add_ticket_images_docx(
    doc: Document,
    image_urls: List[str],
    kb_base_url: Optional[str],
    fetcher: AttachmentFetcher,
    max_width_emu: int,
    max_height_emu: int,
    add_placeholders: bool,
    image_optimize: bool,
    image_target_dpi: int,
    image_max_width_px: Optional[int],
    image_jpeg_quality: int,
    image_convert_png_to_jpeg: bool,
    image_force_jpeg: bool,
    image_min_recompress_bytes: int,
    image_jpeg_optimize: bool,
    image_jpeg_progressive: bool,
    image_png_compress_level: int,
    image_workers: int,
    image_cache: Optional[dict] = None,
    image_dedupe: bool = True,
    image_dedupe_mode: str = "ahash",
    image_dedupe_threshold: int = 5,
) -> None:
    if not image_urls:
        return

    # Per-ticket deduplication state
    seen_exact: set[str] = set()
    seen_ahash: List[int] = []

    # Use a cache to avoid re-fetching/re-optimizing the same URL multiple times within the batch
    cache = image_cache if image_cache is not None else {}

    # Resolve and fetch sequentially (reuse HTTP sessions safely)
    ordered: List[Tuple[str, bool]] = []  # (abs_url, auth_hint)
    fetched: dict[str, Optional[bytes]] = {}

    for raw in image_urls:
        if not raw:
            continue
        clean = sanitize_url(raw)
        if not clean:
            continue
        abs_url = resolve_url(clean, kb_base_url) or clean

        # Track original order and auth hint for placeholders
        auth_hint = False

        # If we already have cached result, just record order
        if abs_url in cache:
            ordered.append((abs_url, auth_hint))
            continue

        data = None

        # Check URL type and prioritize accordingly
        if abs_url.startswith(("http://", "https://")) and is_jira_attachment_url(abs_url):
            data = fetcher.fetch_jira_attachment(abs_url)
            auth_hint = True
        else:
            fid = extract_file_id_from_url(abs_url)
            if fid:
                data = fetcher.fetch_attachment_by_id(fid)
                auth_hint = True
            if not data and abs_url.startswith(("http://", "https://")):
                data = fetcher.fetch_generic_image(abs_url)

        fetched[abs_url] = data if data else None
        ordered.append((abs_url, auth_hint))

    # Compute effective max width in pixels based on doc width and target DPI
    effective_max_width_px = image_max_width_px
    if not effective_max_width_px and image_target_dpi and image_target_dpi > 0:
        effective_max_width_px = int((max_width_emu / EMU_PER_INCH) * image_target_dpi)

    # Optimize in parallel (CPU-bound) to speed up heavy recompression
    to_optimize: List[Tuple[str, bytes]] = []
    for url, data in fetched.items():
        if url in cache:
            continue
        if data:
            if image_optimize:
                to_optimize.append((url, data))
            else:
                cache[url] = data
        else:
            cache[url] = b""  # mark as failed

    if to_optimize:
        workers = max(1, int(image_workers or 1))
        if workers > 1:
            with ThreadPoolExecutor(max_workers=workers) as ex:
                future_map = {
                    ex.submit(
                        optimize_image_bytes,
                        data,
                        max_width_px=effective_max_width_px,
                        jpeg_quality=image_jpeg_quality,
                        convert_png_to_jpeg=image_convert_png_to_jpeg,
                        keep_png_if_transparent=not image_force_jpeg,
                        force_jpeg=image_force_jpeg,
                        force_recompress_min_bytes=image_min_recompress_bytes,
                        jpeg_optimize=image_jpeg_optimize,
                        jpeg_progressive=image_jpeg_progressive,
                        png_compress_level=image_png_compress_level,
                    ): url
                    for (url, data) in to_optimize
                }
                for fut in as_completed(future_map):
                    url = future_map[fut]
                    try:
                        cache[url] = fut.result() or b""
                    except Exception:
                        cache[url] = b""
        else:
            # Sequential fallback
            for url, data in to_optimize:
                try:
                    cache[url] = optimize_image_bytes(
                        data,
                        max_width_px=effective_max_width_px,
                        jpeg_quality=image_jpeg_quality,
                        convert_png_to_jpeg=image_convert_png_to_jpeg,
                        keep_png_if_transparent=not image_force_jpeg,
                        force_jpeg=image_force_jpeg,
                        force_recompress_min_bytes=image_min_recompress_bytes,
                        jpeg_optimize=image_jpeg_optimize,
                        jpeg_progressive=image_jpeg_progressive,
                        png_compress_level=image_png_compress_level,
                    ) or b""
                except Exception:
                    cache[url] = b""

    # Finally, embed in the original order
    for abs_url, auth_hint in ordered:
        data = cache.get(abs_url, b"")
        if not data:
            if add_placeholders:
                add_image_placeholder_docx(doc, abs_url, auth_hint=auth_hint)
            continue

        # Per-ticket duplicate suppression
        if image_dedupe:
            if image_dedupe_mode == "exact":
                try:
                    h = hashlib.blake2b(data, digest_size=16).hexdigest()
                except Exception:
                    h = None
                if h:
                    if h in seen_exact:
                        continue
                    seen_exact.add(h)
            else:
                ah = average_hash(data, 8)
                if ah is not None:
                    if any(hamming_distance(ah, prev) <= image_dedupe_threshold for prev in seen_ahash):
                        continue
                    seen_ahash.append(ah)

        dims = _bytes_to_image_dims_emu(data)
        if not dims:
            if add_placeholders:
                add_image_placeholder_docx(doc, abs_url, auth_hint=auth_hint)
            continue

        w_emu, h_emu = dims
        if w_emu <= 0 or h_emu <= 0:
            if add_placeholders:
                add_image_placeholder_docx(doc, abs_url, auth_hint=auth_hint)
            continue

        scale = min(max_width_emu / w_emu, max_height_emu / h_emu, 1.0)
        target_w = int(w_emu * scale)
        bio = io.BytesIO(data)
        try:
            doc.add_picture(bio, width=Emu(target_w))
        except Exception:
            if add_placeholders:
                add_image_placeholder_docx(doc, abs_url, auth_hint=auth_hint)
            continue


# ---- Ticket rendering ----

def build_doc_for_tickets(
    doc: Document,
    tickets_subset: List[dict],
    env_base: Optional[str],
    fetcher: AttachmentFetcher,
    include_images: bool,
    add_placeholders: bool,
    image_optimize: bool,
    image_target_dpi: int,
    image_max_width_px: Optional[int],
    image_jpeg_quality: int,
    image_convert_png_to_jpeg: bool,
    image_force_jpeg: bool,
    image_min_recompress_bytes: int,
    image_jpeg_optimize: bool,
    image_jpeg_progressive: bool,
    image_png_compress_level: int,
    image_workers: int,
    image_dedupe: bool,
    image_dedupe_mode: str,
    image_dedupe_threshold: int,
) -> None:
    usable_w_emu, usable_h_emu = get_usable_emu(doc)
    image_cache: dict = {}

    for t in tickets_subset:
        subject = (t.get("subject") or "").strip() or "(No subject)"
        problem = (t.get("problem") or "").rstrip()
        solution = (t.get("solution") or "").rstrip()
        image_urls = t.get("image_urls") or []

        # Header
        title_p = doc.add_paragraph(xml_safe(subject))
        title_p.style = "Heading 1"

        # Meta line
        meta_parts = []
        if t.get("ticket_id") is not None:
            meta_parts.append(f"Ticket ID: {t['ticket_id']}")
        if t.get("date"):
            meta_parts.append(str(t["date"]))
        if meta_parts:
            p = add_paragraph_with_inline_formatting(doc, " • ".join(meta_parts))
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        else:
            doc.add_paragraph()  # spacer

        # Problem section
        h = doc.add_paragraph("Problem")
        h.style = "Heading 2"
        add_plain_text_block(doc, problem)

        # Images (after Problem)
        if include_images:
            add_ticket_images_docx(
                doc=doc,
                image_urls=image_urls,
                kb_base_url=env_base,
                fetcher=fetcher,
                max_width_emu=usable_w_emu,
                max_height_emu=usable_h_emu,
                add_placeholders=add_placeholders,
                image_optimize=image_optimize,
                image_target_dpi=image_target_dpi,
                image_max_width_px=image_max_width_px,
                image_jpeg_quality=image_jpeg_quality,
                image_convert_png_to_jpeg=image_convert_png_to_jpeg,
                image_force_jpeg=image_force_jpeg,
                image_min_recompress_bytes=image_min_recompress_bytes,
                image_jpeg_optimize=image_jpeg_optimize,
                image_jpeg_progressive=image_jpeg_progressive,
                image_png_compress_level=image_png_compress_level,
                image_workers=image_workers,
                image_cache=image_cache,
                image_dedupe=image_dedupe,
                image_dedupe_mode=image_dedupe_mode,
                image_dedupe_threshold=image_dedupe_threshold,
            )

        # Solution section
        doc.add_paragraph()  # small spacer
        h2 = doc.add_paragraph("Solution")
        h2.style = "Heading 2"
        add_plain_text_block(doc, solution)


# ---- CLI ----

def main():
    parser = argparse.ArgumentParser(description="Generate DOCX files with configurable number of tickets per file (API-first image fetching).")
    parser.add_argument("--input", "-i", default="Ticket_Data.JSON", help="Path to JSON file with an array of tickets")
    parser.add_argument("--output-dir", "-o", default="documents", help="Output directory for DOCX files")
    parser.add_argument("--tickets-per-file", "-t", type=int, default=50, help="Number of tickets to combine into one DOCX file (default: 50)")
    parser.add_argument("--page-size", choices=["A4", "LETTER"], default="A4", help="Page size")
    parser.add_argument("--margin", type=float, default=36.0, help="Margins in points (default ~0.5 inch)")
    parser.add_argument("--include-images", type=str2bool, default=True, help="Include images from image_urls[]")
    parser.add_argument("--timeout", type=float, default=15.0, help="HTTP timeout seconds for API/generic image downloads")
    parser.add_argument("--image-placeholder", type=str2bool, default=True, help="Insert a textual placeholder when image download fails")
    parser.add_argument("--verbose", type=str2bool, default=False, help="Enable verbose logging")
    parser.add_argument("--base-url", default=None, help="Base URL to resolve relative links and derive API root (overrides JITBIT_BASE_URL)")
    parser.add_argument("--token", default=None, help="Bearer token for Jitbit API (overrides JITBIT_API_TOKEN)")
    parser.add_argument("--jira-token", default=None, help="Bearer token for Jira API (overrides JIRA_API_TOKEN)")
    # Image optimization flags
    parser.add_argument("--image-optimize", type=str2bool, default=True, help="Optimize and compress images before embedding (default: true)")
    parser.add_argument("--image-target-dpi", type=int, default=150, help="Target DPI used to derive max pixel width from document width (default: 150)")
    parser.add_argument("--image-max-width-px", type=int, default=None, help="Explicit max image width in pixels; overrides target-DPI derived width when set")
    parser.add_argument("--image-jpeg-quality", type=int, default=75, help="JPEG quality for recompression (default: 75)")
    parser.add_argument("--image-convert-png-to-jpeg", type=str2bool, default=True, help="Convert PNG to JPEG when transparency is not required (default: true)")
    parser.add_argument("--image-force-jpeg", "-image-force-jpeg", type=str2bool, default=False, help="Force JPEG conversion even if image has transparency (default: false)")
    parser.add_argument("--image-min-recompress-bytes", type=int, default=131072, help="Only recompress JPEG/PNG if original size is at least this many bytes (default: 131072)")
    parser.add_argument("--image-jpeg-optimize", type=str2bool, default=False, help="Enable extra JPEG encoder optimization (slower). Default: false")
    parser.add_argument("--image-jpeg-progressive", type=str2bool, default=False, help="Encode progressive JPEGs (slower to encode). Default: false")
    parser.add_argument("--image-png-compress-level", type=int, default=6, help="PNG zlib compression level 0-9 (higher = smaller but slower). Default: 6")
    parser.add_argument("--image-workers", type=int, default=0, help="Parallel workers for image optimization (0=auto)")
    parser.add_argument("--image-dedupe", type=str2bool, default=True, help="Deduplicate similar images within a ticket (default: true)")
    parser.add_argument("--image-dedupe-mode", choices=["exact", "ahash"], default="ahash", help="Deduplication mode: exact byte hash or perceptual average-hash")
    parser.add_argument("--image-dedupe-threshold", type=int, default=5, help="Hamming distance threshold for ahash mode (default: 5)")
    args = parser.parse_args()

    # Compute default workers if needed
    if args.image_workers is None or args.image_workers <= 0:
        try:
            cpu = os.cpu_count() or 4
        except Exception:
            cpu = 4
        args.image_workers = min(32, cpu * 2)

    # Load JSON (top-level array)
    with open(args.input, "r", encoding="utf-8") as f:
        data = json.load(f)

    if isinstance(data, dict):
        tickets = data.get("tickets") or data.get("data") or []
        if not isinstance(tickets, list):
            print("[WARN] Input JSON is an object but no 'tickets' array found. Expecting a top-level array.", file=sys.stderr)
            tickets = []
    elif isinstance(data, list):
        tickets = data
    else:
        print("[WARN] Input JSON not an array.", file=sys.stderr)
        tickets = []

    env_base = (args.base_url or os.getenv("JITBIT_BASE_URL", "") or "").strip() or None
    api_root = derive_api_root(None, env_base, verbose=args.verbose)
    token = (args.token or os.getenv("JITBIT_API_TOKEN", "") or "").strip()
    jira_token = (args.jira_token or os.getenv("JIRA_API_TOKEN", "") or "").strip()
    jira_email = (os.getenv("JIRA_EMAIL", "") or "").strip()

    if not env_base:
        print("[WARN] JITBIT_BASE_URL not set in environment/.env. Relative Jitbit links cannot be resolved.", file=sys.stderr)
    if not token:
        print("[WARN] JITBIT_API_TOKEN not set in environment/.env. Jitbit-protected images will not load.", file=sys.stderr)
    if not jira_token:
        print("[WARN] JIRA_API_TOKEN not set in environment/.env. Jira attachment downloads may fail.", file=sys.stderr)
    if not jira_email:
        print("[WARN] JIRA_EMAIL not set in environment/.env. Jira Basic auth is not configured; use JIRA_EMAIL + JIRA_API_TOKEN.", file=sys.stderr)

    fetcher = AttachmentFetcher(api_root=api_root, token=token, jira_token=jira_token, jira_email=jira_email, timeout=args.timeout, verbose=args.verbose)

    # Create output directory if it doesn't exist
    os.makedirs(args.output_dir, exist_ok=True)

    # Validate tickets_per_file parameter
    if args.tickets_per_file < 1:
        print("[ERROR] --tickets-per-file must be at least 1", file=sys.stderr)
        sys.exit(1)

    # Process tickets in batches
    total_tickets = len(tickets)
    if total_tickets == 0:
        print("[INFO] No tickets found in input file.")
        return

    total_batches = math.ceil(total_tickets / args.tickets_per_file)
    t0 = time.time()

    batch_count = 0
    generated_files = 0

    for i in range(0, total_tickets, args.tickets_per_file):
        batch_count += 1
        ticket_batch = tickets[i:i + args.tickets_per_file]
        batch_start = time.time()
        
        # Create filename for this batch
        start_idx = i + 1
        end_idx = min(i + args.tickets_per_file, total_tickets)
        
        if args.tickets_per_file == 1:
            # Special case: one ticket per file, use original naming pattern
            ticket = ticket_batch[0]
            ticket_id = ticket.get("ticket_id", "unknown")
            subject = ticket.get("subject", "No Subject").strip()
            safe_subject = re.sub(r'[<>:"/\\|?*]', '_', subject)[:50]
            filename = f"ticket_{ticket_id}_{safe_subject}.docx"
        else:
            # Multiple tickets per file
            filename = f"tickets_{start_idx:04d}-{end_idx:04d}_batch_{batch_count:03d}.docx"
        
        out_file = os.path.join(args.output_dir, filename)

        # Create a new document for this batch
        doc = Document()
        set_page_size_and_margins(doc, args.page_size, args.margin)

        # Build document with this batch of tickets
        build_doc_for_tickets(
            doc=doc,
            tickets_subset=ticket_batch,
            env_base=env_base,
            fetcher=fetcher,
            include_images=args.include_images,
            add_placeholders=args.image_placeholder,
            image_optimize=args.image_optimize,
            image_target_dpi=args.image_target_dpi,
            image_max_width_px=args.image_max_width_px,
            image_jpeg_quality=args.image_jpeg_quality,
            image_convert_png_to_jpeg=args.image_convert_png_to_jpeg,
            image_force_jpeg=args.image_force_jpeg,
            image_min_recompress_bytes=args.image_min_recompress_bytes,
            image_jpeg_optimize=args.image_jpeg_optimize,
            image_jpeg_progressive=args.image_jpeg_progressive,
            image_png_compress_level=args.image_png_compress_level,
            image_workers=args.image_workers,
            image_dedupe=args.image_dedupe,
            image_dedupe_mode=args.image_dedupe_mode,
            image_dedupe_threshold=args.image_dedupe_threshold,
        )
        
        doc.save(out_file)
        generated_files += 1

        batch_elapsed = time.time() - batch_start
        elapsed_total = time.time() - t0
        batches_done = generated_files
        remaining = max(0, total_batches - batches_done)
        if batches_done > 0 and remaining > 0:
            avg = elapsed_total / batches_done
            eta_seconds = int(avg * remaining)
            eta_h = eta_seconds // 3600
            eta_m = (eta_seconds % 3600) // 60
            eta_s = eta_seconds % 60
            eta_hms = f"{eta_h:d}:{eta_m:02d}:{eta_s:02d}"
            eta_clock = datetime.now() + timedelta(seconds=eta_seconds)
            extra = f" | time: {batch_elapsed:.1f}s | ETA: {eta_hms} (~{eta_clock.strftime('%H:%M:%S')})"
        else:
            extra = f" | time: {batch_elapsed:.1f}s"

        print(f"[OK] DOCX generated: {out_file} (tickets {start_idx}-{end_idx}, {len(ticket_batch)} tickets){extra}")

    if args.tickets_per_file == 1:
        print(f"[INFO] Generated {generated_files} separate DOCX files in {args.output_dir}/")
    else:
        print(f"[INFO] Generated {generated_files} DOCX files containing {total_tickets} tickets ({args.tickets_per_file} tickets per file) in {args.output_dir}/")


if __name__ == "__main__":
    main()
