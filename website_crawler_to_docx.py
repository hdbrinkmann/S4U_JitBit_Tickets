#!/usr/bin/env python3
"""
Website Crawler to Word Document Generator

This script crawls a website starting from a given URL, follows all internal links,
extracts content from each page, and generates a well-formatted Word document
with all the collected content.
"""

import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
import logging
import argparse
from collections import deque
import re
from datetime import datetime


class WebsiteCrawler:
    def __init__(self, start_url, max_depth=3, delay=1):
        """
        Initialize the website crawler.
        
        Args:
            start_url (str): The starting URL to crawl
            max_depth (int): Maximum depth to crawl (default: 3)
            delay (float): Delay between requests in seconds (default: 1)
        """
        self.start_url = start_url
        self.base_domain = urlparse(start_url).netloc
        self.max_depth = max_depth
        self.delay = delay
        self.visited_urls = set()
        self.pages_content = []
        
        # Setup logging
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )
        self.logger = logging.getLogger(__name__)
        
        # Setup session with headers
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })

    def is_valid_url(self, url):
        """Check if URL is valid and belongs to the same domain."""
        try:
            parsed = urlparse(url)
            return (
                parsed.netloc == self.base_domain and
                parsed.scheme in ['http', 'https'] and
                not any(ext in url.lower() for ext in ['.pdf', '.jpg', '.png', '.gif', '.zip', '.exe'])
            )
        except:
            return False

    def extract_links(self, soup, current_url):
        """Extract all valid links from the current page."""
        links = set()
        for link in soup.find_all('a', href=True):
            url = urljoin(current_url, link['href'])
            if self.is_valid_url(url):
                links.add(url)
        return links

    def clean_text(self, text):
        """Clean and normalize text content."""
        # Remove extra whitespace and normalize
        text = re.sub(r'\s+', ' ', text.strip())
        return text

    def extract_content(self, soup, url):
        """Extract meaningful content from the page."""
        # Remove script and style elements
        for script in soup(["script", "style", "nav", "footer", "header"]):
            script.decompose()

        # Get page title - try to extract API-specific title
        title_text = self.extract_api_title(soup, url)

        # Get main content
        content_selectors = [
            'main', 'article', '.content', '#content', '.main-content',
            '.post-content', 'div[role="main"]'
        ]
        
        main_content = None
        for selector in content_selectors:
            main_content = soup.select_one(selector)
            if main_content:
                break
        
        if not main_content:
            main_content = soup.find('body')

        # Extract text content
        if main_content:
            # Get all paragraphs, headers, lists, etc.
            content_elements = []
            
            for element in main_content.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li', 'div']):
                text = self.clean_text(element.get_text())
                if text and len(text) > 10:  # Only include substantial text
                    element_type = element.name
                    content_elements.append({
                        'type': element_type,
                        'text': text,
                        'level': int(element_type[1]) if element_type.startswith('h') else 0
                    })

            return {
                'url': url,
                'title': title_text,
                'content_elements': content_elements
            }
        
        return {
            'url': url,
            'title': title_text,
            'content_elements': []
        }

    def extract_api_title(self, soup, url):
        """Extract a meaningful title for API endpoints."""
        # Check if this is an API operation page
        if '?op=' in url:
            # Extract operation name from URL
            operation = url.split('?op=')[-1]
            operation_clean = operation.replace('_', ' ').title()
            
            # Try to find a more descriptive title in the page content
            h1_elements = soup.find_all('h1')
            h2_elements = soup.find_all('h2')
            
            for heading in h1_elements + h2_elements:
                heading_text = heading.get_text().strip()
                if heading_text and len(heading_text) > 3 and heading_text != "TimeMapAPI":
                    return f"{operation_clean} - {heading_text}"
            
            # Look for specific patterns that might indicate the API function
            paragraphs = soup.find_all('p')
            for p in paragraphs:
                text = p.get_text().strip()
                if text.startswith('The following operations are supported'):
                    continue
                if len(text) > 20 and len(text) < 200:
                    # This might be a description
                    return f"{operation_clean} API"
            
            return f"{operation_clean} API Endpoint"
        
        elif 'WSDL' in url:
            return "WSDL Service Description"
        
        # Default title extraction
        title = soup.find('title')
        if title:
            title_text = title.get_text().strip()
            if title_text and title_text != "TimeMapAPI":
                return title_text
        
        # Try to extract from main heading
        h1 = soup.find('h1')
        if h1:
            h1_text = h1.get_text().strip()
            if h1_text and h1_text != "TimeMapAPI":
                return h1_text
        
        return "API Documentation"

    def crawl_page(self, url, depth=0):
        """Crawl a single page and return its content and links."""
        if url in self.visited_urls or depth > self.max_depth:
            return set()

        try:
            self.logger.info(f"Crawling (depth {depth}): {url}")
            response = self.session.get(url, timeout=30)
            response.raise_for_status()
            
            self.visited_urls.add(url)
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Extract content
            page_content = self.extract_content(soup, url)
            self.pages_content.append(page_content)
            
            # Extract links for further crawling
            links = self.extract_links(soup, url)
            
            # Add delay to be respectful
            time.sleep(self.delay)
            
            return links

        except Exception as e:
            self.logger.error(f"Error crawling {url}: {e}")
            return set()

    def crawl_website(self):
        """Crawl the entire website using breadth-first search."""
        self.logger.info(f"Starting to crawl website: {self.start_url}")
        
        # Queue for BFS: (url, depth)
        queue = deque([(self.start_url, 0)])
        
        while queue:
            current_url, depth = queue.popleft()
            
            if current_url not in self.visited_urls and depth <= self.max_depth:
                links = self.crawl_page(current_url, depth)
                
                # Add new links to queue
                for link in links:
                    if link not in self.visited_urls:
                        queue.append((link, depth + 1))
        
        self.logger.info(f"Crawling completed. Visited {len(self.visited_urls)} pages.")

    def create_word_document(self, output_filename):
        """Create a well-formatted Word document with the crawled content."""
        self.logger.info("Creating Word document...")
        
        doc = Document()
        
        # Set up document styles
        self.setup_document_styles(doc)
        
        # Add title
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(f"Website Crawl Report: {self.base_domain}")
        title_run.font.size = Pt(18)
        title_run.bold = True
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Starting URL: {self.start_url}")
        doc.add_paragraph(f"Total pages crawled: {len(self.pages_content)}")
        doc.add_paragraph("")
        
        # Add table of contents
        doc.add_heading("Table of Contents", level=1)
        for i, page in enumerate(self.pages_content, 1):
            toc_paragraph = doc.add_paragraph(f"{i}. {page['title']}")
            toc_paragraph.style = 'List Number'
        doc.add_page_break()
        
        # Add content for each page as chapters
        for i, page in enumerate(self.pages_content, 1):
            # Chapter header (no "Page X:" prefix)
            doc.add_heading(page['title'], level=1)
            
            # Add URL as subtitle in smaller text
            url_paragraph = doc.add_paragraph()
            url_run = url_paragraph.add_run(f"Endpoint: {page['url']}")
            url_run.font.size = Pt(9)
            url_run.italic = True
            doc.add_paragraph("")
            
            # Page content
            for element in page['content_elements']:
                if element['type'].startswith('h') and element['level'] > 0:
                    # Header
                    doc.add_heading(element['text'], level=min(element['level'] + 1, 6))
                elif element['type'] == 'p':
                    # Paragraph
                    doc.add_paragraph(element['text'])
                elif element['type'] == 'li':
                    # List item
                    doc.add_paragraph(element['text'], style='List Bullet')
                else:
                    # Other content
                    doc.add_paragraph(element['text'])
            
            # Add page break between chapters (except for the last one)
            if i < len(self.pages_content):
                doc.add_page_break()
        
        # Save document
        doc.save(output_filename)
        self.logger.info(f"Word document saved as: {output_filename}")

    def setup_document_styles(self, doc):
        """Setup custom styles for the document."""
        styles = doc.styles
        
        # Create a custom normal style if needed
        try:
            normal_style = styles['Normal']
            normal_style.font.size = Pt(11)
            normal_style.font.name = 'Calibri'
        except:
            pass


def main():
    """Main function to run the website crawler."""
    parser = argparse.ArgumentParser(description='Crawl a website and generate a Word document')
    parser.add_argument('url', help='The starting URL to crawl')
    parser.add_argument('-o', '--output', default='website_crawl_report.docx',
                       help='Output Word document filename (default: website_crawl_report.docx)')
    parser.add_argument('-d', '--depth', type=int, default=3,
                       help='Maximum crawl depth (default: 3)')
    parser.add_argument('--delay', type=float, default=1.0,
                       help='Delay between requests in seconds (default: 1.0)')
    
    args = parser.parse_args()
    
    # Create crawler instance
    crawler = WebsiteCrawler(
        start_url=args.url,
        max_depth=args.depth,
        delay=args.delay
    )
    
    try:
        # Crawl the website
        crawler.crawl_website()
        
        # Create Word document
        crawler.create_word_document(args.output)
        
        print(f"\n✅ Successfully created Word document: {args.output}")
        print(f"📊 Total pages crawled: {len(crawler.pages_content)}")
        print(f"🔗 Total URLs visited: {len(crawler.visited_urls)}")
        
    except KeyboardInterrupt:
        print("\n⚠️  Crawling interrupted by user")
    except Exception as e:
        print(f"\n❌ Error: {e}")


if __name__ == "__main__":
    main()
